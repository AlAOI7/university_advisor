import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_path = r'C:\Users\ALAOI\university_advisor\Chapter_3_English.docx'

doc = docx.Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ─────────────────────────────────────────
# Chapter 3: System Design
# ─────────────────────────────────────────
add_heading('Chapter 3', 1)
add_heading('System Design', 2)
add_para(
    "This chapter provides a comprehensive and detailed presentation of the University Advisor System's "
    "design across all four major components: the Front-End, the Back-End, the Artificial Intelligence (AI) "
    "module, and the administrative Dashboard. The system design phase translates the requirements gathered "
    "during analysis into a concrete blueprint for implementation. Each section describes the architectural "
    "decisions, interface designs, data models, and intelligent components that together form a cohesive, "
    "scalable, and user-friendly academic advising platform."
)

# ─────────────────────────────────────────
# 3.1 Front-End Section
# ─────────────────────────────────────────
add_heading('3.1 Front-End Section', 2)
add_para(
    "The front-end of the University Advisor System is engineered to deliver an intuitive, modern, and "
    "fully responsive user experience. The design philosophy is centered around clarity and accessibility, "
    "ensuring that students, authors, editors, and administrators can interact with the system effortlessly. "
    "The interface is built with HTML5, CSS3, JavaScript, and Tailwind CSS, and adapts seamlessly to all "
    "screen sizes and device types. Each page and view is described in detail in the following subsections."
)

# 3.1.1
add_heading('3.1.1 Homepage – Main Entry Point', 3)
add_para(
    "The Homepage serves as the primary landing page for all users visiting the University Advisor System. "
    "It presents a visually compelling introduction to the platform, featuring a hero section with a clear "
    "call-to-action that guides visitors to either register for a new account or log in to an existing one. "
    "The page highlights the system's core features, including the academic advising engine, the scientific "
    "journal, and the AI-powered classification tool. A clean and modern navigation bar provides quick "
    "access to all major sections of the platform. The homepage is fully responsive, ensuring an optimal "
    "viewing experience on desktop, tablet, and mobile devices."
)

# 3.1.2
add_heading('3.1.2 Scientific Journal – Main Page', 3)
add_para(
    "The Scientific Journal Main Page acts as the central hub for all published academic content within "
    "the system. It displays a structured list of journal issues and articles, each accompanied by its "
    "title, abstract snippet, author names, keywords, and publication date. Users can browse through "
    "articles by issue, filter by keyword or category, and search for specific topics of interest. "
    "The layout is designed for academic readability, with a clean typographic hierarchy that makes "
    "navigating through research content straightforward. Pagination controls allow users to move through "
    "large volumes of content efficiently."
)

# 3.1.3
add_heading('3.1.3 Authentication – Login View', 3)
add_para(
    "The Authentication Login View provides a secure and straightforward interface for users to access "
    "their accounts. The form requests the user's registered email address and password. Client-side "
    "validation provides immediate feedback on input errors, while server-side authentication ensures "
    "that credentials are verified securely before granting access. The page also features a 'Forgot "
    "Password' link to initiate the account recovery process and a link to the registration page for "
    "new users. Upon successful login, users are redirected to their role-specific dashboard, whether "
    "they are a student, author, editor, or administrator."
)

# 3.1.4
add_heading('3.1.4 Scientific Journal – Article Details Page', 3)
add_para(
    "The Article Details Page presents a comprehensive view of a single research article. It displays "
    "the full title, the complete list of authors and their affiliations, the publication date and issue "
    "number, the full abstract, and a list of keywords. Where available, a link to the full PDF document "
    "is provided for download. The page also lists the article's references in a standardized academic "
    "format. Related articles are suggested at the bottom of the page to encourage further exploration "
    "of the journal's content. The page is optimized for readability with appropriate typography and "
    "spacing."
)

# 3.1.5
add_heading('3.1.5 Editor Access – Dashboard Visibility', 3)
add_para(
    "This view describes the portal that editors see upon logging into the system. Based on the editor's "
    "role and assigned permissions, the navigation menu dynamically updates to display editor-specific "
    "options, such as the Article Review Panel and Issue Management. The main dashboard surface provides "
    "a summary of pending articles awaiting review, recently submitted manuscripts, and any notifications "
    "or messages from authors. This role-based visibility ensures that the interface is uncluttered and "
    "relevant to the editor's specific responsibilities within the system."
)

# 3.1.6
add_heading('3.1.6 Editor Dashboard – Article Review Panel', 3)
add_para(
    "The Article Review Panel is a dedicated workspace for editors to manage the peer-review process. "
    "It presents a tabular list of all articles assigned to the editor, categorized by their current "
    "review status: Pending, Under Review, Revision Requested, Accepted, or Rejected. Each row in the "
    "table displays the article title, the submitting author, the submission date, and the current "
    "status. Editors can filter and sort articles by any of these criteria to efficiently manage their "
    "workload. Action buttons allow the editor to open an article for detailed review, assign it to a "
    "specific reviewer, or update its status directly from the panel."
)

# 3.1.7
add_heading('3.1.7 Editor Dashboard – Article Review Interface', 3)
add_para(
    "The Article Review Interface provides editors with all the information and tools needed to evaluate "
    "a single manuscript. The full article details, including the abstract, content, figures, and "
    "references, are displayed for review. The editor can view the manuscript's metadata and any "
    "supplementary files provided by the author. A dedicated review form allows the editor to record "
    "their evaluation, select a recommendation (Accept, Minor Revision, Major Revision, Reject), and "
    "write detailed comments for the author. The submitted review is automatically timestamped and "
    "logged in the system, and the author is notified of the decision via the platform."
)

# 3.1.8
add_heading('3.1.8 Author Access – Dashboard Visibility', 3)
add_para(
    "Similar to the editor access view, the Author's dashboard visibility is tailored specifically to "
    "the needs of a researcher submitting work to the journal. Upon login, an author is greeted with a "
    "personalized dashboard that shows a summary of their submitted articles and their current statuses. "
    "The navigation provides access to the article submission form, their profile, and any notifications "
    "from editors regarding their submissions. This role-specific view simplifies the author's experience "
    "by presenting only the tools and information that are directly relevant to the publication workflow."
)

# 3.1.9
add_heading('3.1.9 Author Dashboard – Article Management Interface', 3)
add_para(
    "The Author's Article Management Interface gives researchers a complete overview of all their "
    "submitted manuscripts. Each article is listed with its title, submission date, current status in "
    "the review pipeline, and any editor feedback. Authors can click on any article to view the full "
    "submission details, read editor comments, download review reports, and submit revised versions "
    "when a revision is requested. The interface is designed to keep the author fully informed about "
    "the progress of their submissions throughout the entire editorial workflow, from initial submission "
    "to final publication."
)

# 3.1.10
add_heading('3.1.10 Author Dashboard – Submit Article Form', 3)
add_para(
    "The Submit Article Form is the interface through which authors contribute new manuscripts to the "
    "journal. The form is structured in a logical, step-by-step format to guide authors through the "
    "submission process. Authors are required to enter the article title, provide a full abstract, "
    "list all co-authors, add relevant keywords, specify the research domain, and upload the manuscript "
    "file in the accepted format (e.g., PDF or DOCX). Fields for optional supplementary materials, "
    "such as datasets or figures, are also provided. Upon submission, the system sends a confirmation "
    "notification to the author and routes the article to the appropriate editor for initial screening."
)

# 3.1.11
add_heading('3.1.11 User Profile – Account Information View', 3)
add_para(
    "The User Profile page allows all registered users to view and manage their personal account "
    "information. Users can update their display name, email address, affiliated institution, and "
    "contact details. A profile picture upload option is provided to personalize the account. "
    "Authors can additionally manage their professional biography, list of research interests, and "
    "ORCID identifier. A password change section is available for security management. All changes "
    "are saved upon confirmation, with appropriate validation to ensure data integrity and security "
    "across all user types."
)

# 3.1.12
add_heading('3.1.12 AI Classification – Landing & Instructions Interface', 3)
add_para(
    "The AI Classification Landing Interface serves as the entry point to the system's AI-powered "
    "image classification module. This page clearly explains the tool's purpose: to classify research "
    "images or datasets using a trained deep learning model. It provides step-by-step instructions "
    "guiding the user through the classification process: first, uploading an image or data file; "
    "second, initiating the prediction; and third, viewing the results. The interface is designed "
    "to be approachable for users who may not have a technical background in AI, with plain-language "
    "descriptions and illustrative icons for each step."
)

# 3.1.13
add_heading('3.1.13 AI Classification – Image Upload Interface', 3)
add_para(
    "The Image Upload Interface is the functional core of the AI classification module. It features "
    "a clear and user-friendly drag-and-drop area for uploading images, with a fallback 'Browse Files' "
    "button for traditional file selection. The interface displays a preview of the uploaded image "
    "immediately after selection, allowing the user to confirm they have chosen the correct file. "
    "Accepted file formats (e.g., JPG, PNG) and maximum file size limits are clearly indicated. Once "
    "the user is satisfied with their selection, a prominently placed 'Classify' button submits the "
    "image to the AI model for processing."
)

# 3.1.14
add_heading('3.1.14 AI Classification – Prediction Result Interface', 3)
add_para(
    "The Prediction Result Interface displays the output of the AI classification model in a clear "
    "and informative manner. After the model processes the uploaded image, the top predicted class "
    "is displayed prominently, accompanied by a confidence score expressed as a percentage. A bar "
    "chart or similar visualization shows the probability distribution across all possible categories, "
    "providing the user with a transparent view of the model's reasoning. The original uploaded image "
    "is displayed alongside the results for reference. Users can choose to classify a new image or "
    "save the results for their records."
)

# 3.1.15
add_heading('3.1.15 Dashboard – Sales by Category', 3)
add_para(
    "The 'Sales by Category' dashboard view provides administrators with a visual breakdown of "
    "platform engagement and activity metrics, organized by category. In the context of the University "
    "Advisor system, this translates to metrics such as the number of article submissions, user "
    "registrations, or test completions, segmented by academic field or major category. An interactive "
    "bar or pie chart allows administrators to quickly identify which categories are most active. "
    "Filtering options by date range and category type enable detailed temporal and comparative analysis, "
    "supporting data-driven decision-making."
)

# 3.1.16
add_heading('3.1.16 Dashboard – Treemap by Region > Category > Sub-Category', 3)
add_para(
    "The Treemap Dashboard View offers a hierarchical visualization of platform data, organized in a "
    "nested structure from Region down to Category and Sub-Category. This powerful visualization tool "
    "allows administrators to spot patterns and proportions at a glance, with larger rectangles "
    "representing higher volumes of activity. For example, an administrator could see which geographic "
    "regions have the most users, and within each region, which academic majors are most popular. "
    "The interactive treemap allows users to click into any segment to drill down further into the data, "
    "making it an effective tool for high-level strategic analysis."
)

# ─────────────────────────────────────────
# 3.2 Back-End Section
# ─────────────────────────────────────────
add_heading('3.2 Back-End Section', 2)
add_para(
    "The back-end architecture of the University Advisor System is built on the Django web framework "
    "using Python, adhering to the Model-View-Template (MVT) design pattern. PostgreSQL serves as "
    "the primary relational database, ensuring data integrity, reliability, and support for complex "
    "queries. The back-end is responsible for all business logic, user authentication, RESTful API "
    "endpoints, and database interactions. The following subsections detail the core data models "
    "that underpin the system's data layer."
)

# 3.2.2 Model Details
add_heading('3.2.2 Model Details', 3)
add_para(
    "The system's database schema is composed of a set of interconnected models that represent the "
    "core entities of the platform. Each model is implemented as a Django ORM class, which maps "
    "directly to a table in the PostgreSQL database. The models are designed to be normalized, "
    "efficient, and reflective of the real-world relationships between entities in an academic "
    "publishing and advising system."
)

# 3.2.2.1 User Model
add_heading('3.2.2.1 User Model', 4)
add_para(
    "The User Model is the foundational entity of the system, extending Django's built-in "
    "AbstractUser model. It stores essential authentication credentials, including a unique "
    "username, email address, and hashed password. It also includes profile fields such as "
    "the user's full name, affiliated institution, and role (e.g., Student, Author, Editor, "
    "Administrator). The role field is critical for enforcing role-based access control (RBAC) "
    "throughout the application, ensuring that each user can only access functionalities "
    "appropriate to their designation."
)

# 3.2.2.2 Journal Model
add_heading('3.2.2.2 Journal Model', 4)
add_para(
    "The Journal Model represents a scientific journal hosted on the platform. It contains "
    "attributes such as the journal's title, ISSN (International Standard Serial Number), "
    "a description of the journal's scope and focus, the founding date, the editor-in-chief "
    "(a ForeignKey relationship to the User Model), and the journal's current status (Active/Archived). "
    "This model serves as the top-level container for organizing all published issues and articles "
    "within the system."
)

# 3.2.2.3 Issue Model
add_heading('3.2.2.3 Issue Model', 4)
add_para(
    "The Issue Model represents a specific volume and issue of a journal. It is related to the "
    "Journal Model through a ForeignKey, establishing a one-to-many relationship where a single "
    "journal can have multiple issues. Key attributes include the volume number, issue number, "
    "publication date, and a description or theme for the issue. The Issue Model serves as the "
    "organizational unit that groups individual articles published together in a single release."
)

# 3.2.2.4 Author Model
add_heading('3.2.2.4 Author Model', 4)
add_para(
    "The Author Model extends the core user profile to capture information specific to researchers "
    "and academics. It maintains a OneToOne relationship with the User Model and stores additional "
    "details such as the author's professional biography, ORCID identifier, primary research "
    "interests, and a list of their institutional affiliations. This dedicated model allows the "
    "system to manage author-specific data separately from general user data, keeping the schema "
    "clean and well-organized."
)

# 3.2.2.5 Editor Model
add_heading('3.2.2.5 Editor Model', 4)
add_para(
    "Similar to the Author Model, the Editor Model maintains a OneToOne relationship with the "
    "User Model and stores data specific to journal editors. This includes their area of expertise, "
    "which is used by the system to intelligently assign submitted articles for review. It also "
    "tracks the journals and issues the editor is responsible for managing. The Editor Model "
    "is central to the workflow routing logic in the back-end."
)

# 3.2.2.6 Article Model
add_heading('3.2.2.6 Article Model', 4)
add_para(
    "The Article Model is one of the most central entities in the system. It stores all information "
    "pertaining to a submitted manuscript, including the title, abstract, the full content or "
    "a link to the uploaded manuscript file, submission date, and the associated journal issue "
    "(ForeignKey to the Issue Model). It maintains a ManyToMany relationship with the Author "
    "Model to support multi-author publications, and includes a status field that tracks the "
    "article through the editorial workflow (e.g., Submitted, Under Review, Accepted, Published, Rejected)."
)

# 3.2.2.7 Review Model
add_heading('3.2.2.7 Review Model', 4)
add_para(
    "The Review Model captures the outcome of an editorial review for a specific article. It "
    "holds ForeignKey relationships to both the Article Model and the Editor Model, identifying "
    "which editor reviewed which article. The model stores the editor's written comments, their "
    "overall recommendation, and the timestamp of the review. This model forms the backbone of "
    "the peer-review audit trail, ensuring that all editorial decisions are recorded and "
    "traceable within the system."
)

# 3.2.2.8 Keyword Model
add_heading('3.2.2.8 Keyword Model', 4)
add_para(
    "The Keyword Model is a simple but essential component that stores individual academic "
    "keywords or terms used to categorize and tag articles. Each keyword is stored as a "
    "unique string entry. Keywords are associated with articles through the ArticleKeywords "
    "junction model, enabling efficient search and filtering of the journal's content by topic. "
    "The use of a dedicated model ensures that keywords are standardized and avoids "
    "duplication across the database."
)

# 3.2.2.9 ArticleKeywords Model
add_heading('3.2.2.9 ArticleKeywords Model', 4)
add_para(
    "The ArticleKeywords Model is a junction table that implements the ManyToMany relationship "
    "between the Article Model and the Keyword Model. Each record in this table links one "
    "article to one keyword, allowing a single article to have multiple keywords and a single "
    "keyword to be associated with multiple articles. This standard relational design pattern "
    "ensures database normalization while providing the flexibility needed for comprehensive "
    "content categorization and search functionality."
)

# 3.2.2.10 Reference Model
add_heading('3.2.2.10 Reference Model', 4)
add_para(
    "The Reference Model stores the bibliographic references cited within an article. It "
    "maintains a ForeignKey relationship to the Article Model, linking each reference to its "
    "parent article. Fields include the full citation text formatted according to the journal's "
    "preferred style (e.g., APA, IEEE), the DOI (Digital Object Identifier) if available, "
    "and the order in which the reference appears in the article. This model enables the system "
    "to display well-formatted reference lists on each article's detail page."
)

# 3.2.3 Key Points
add_heading('3.2.3 Key Points', 3)
add_para(
    "The back-end architecture is guided by several core design principles that ensure the "
    "system's quality and maintainability:"
)
add_bullet("Separation of Concerns: Business logic is encapsulated within Django's service layer and model methods, keeping views thin and focused on request/response handling.")
add_bullet("Security First: All endpoints are protected by Django's CSRF middleware, and user authentication is enforced on all private routes using Django's built-in permission system.")
add_bullet("RESTful API Design: The API endpoints follow REST conventions, using appropriate HTTP methods (GET, POST, PUT, DELETE) and returning standardized JSON responses.")
add_bullet("Scalability: The use of PostgreSQL and a normalized schema ensures the system can handle a growing volume of users and articles without performance degradation.")
add_bullet("Data Integrity: ForeignKey and ManyToMany relationships are enforced at the database level, ensuring that referential integrity is maintained at all times.")

# 3.2.5 ERD
add_heading('3.2.5 Entity Relationship Diagram (ERD)', 3)
add_para(
    "The Entity Relationship Diagram (ERD) provides a comprehensive visual representation of the "
    "database schema and the relationships between all models in the University Advisor System. "
    "The diagram illustrates both the cardinality (one-to-one, one-to-many, many-to-many) and "
    "the nature of the relationships between entities."
)
add_para(
    "Key relationships illustrated in the ERD include: the User model as a central hub connected "
    "to Author and Editor through one-to-one extensions; the Journal connected to Issues in a "
    "one-to-many relationship; Issues connected to Articles in a one-to-many relationship; "
    "Articles connected to Authors through the many-to-many relationship; Articles connected "
    "to Keywords through the ArticleKeywords junction table; and Articles connected to Reviews "
    "and References through one-to-many relationships. This diagram serves as the authoritative "
    "reference for the system's data structure and was used as the primary guide during "
    "the database implementation phase."
)

# ─────────────────────────────────────────
# 3.3 AI Section
# ─────────────────────────────────────────
add_heading('3.3 AI Section', 2)
add_para(
    "The Artificial Intelligence section of the University Advisor System is designed to provide "
    "intelligent, data-driven functionalities that elevate the platform beyond a static information "
    "repository. The AI pipeline encompasses four primary stages: data ingestion and preprocessing, "
    "model architecture definition, training, and evaluation. The primary AI feature is an image "
    "classification model that can categorize research-related images, supporting the journal's "
    "editorial and student exploration workflows."
)

# 3.3.1
add_heading('3.3.1 Data Ingestion and Preprocessing', 3)
add_para(
    "The first stage of the AI pipeline involves collecting and preparing the data required to "
    "train the classification model. Raw image datasets are ingested from designated data sources, "
    "which may include curated academic image databases or custom datasets assembled for this project. "
    "The preprocessing pipeline performs several operations to standardize the data: images are "
    "resized to a uniform resolution required by the model (e.g., 224x224 pixels), pixel values "
    "are normalized to a [0, 1] range to improve training stability, and data augmentation "
    "techniques such as random horizontal flipping, rotation, and color jitter are applied to "
    "artificially expand the training set and improve the model's ability to generalize to "
    "unseen data. The dataset is then split into training, validation, and test sets using a "
    "stratified approach to ensure balanced class representation in each split."
)

# 3.3.2
add_heading('3.3.2 Model Architecture', 3)
add_para(
    "The classification model is built upon a Convolutional Neural Network (CNN) architecture, "
    "chosen for its proven effectiveness in image recognition tasks. To leverage the power of "
    "transfer learning and reduce training time, the system uses a pre-trained backbone such as "
    "ResNet-50 or EfficientNet, which has been pre-trained on the large-scale ImageNet dataset. "
    "The pre-trained convolutional base is used as a feature extractor, with its weights initially "
    "frozen. A custom classification head is appended on top of the backbone, consisting of a "
    "Global Average Pooling layer, followed by one or more fully connected (Dense) layers with "
    "ReLU activation, Dropout layers for regularization, and a final Softmax output layer with "
    "neurons corresponding to the number of target classes. The model is implemented using "
    "PyTorch, allowing for flexible and efficient model definition and training."
)

# 3.3.3
add_heading('3.3.3 Training Engine', 3)
add_para(
    "The training engine orchestrates the model training process. The Adam optimizer is used "
    "with an initial learning rate of 1e-4, chosen for its adaptive learning rate properties "
    "which generally lead to faster convergence. Cross-Entropy Loss is used as the loss function, "
    "as it is well-suited for multi-class classification problems. Training is conducted in two "
    "phases: first, only the classification head weights are updated with the backbone frozen "
    "(feature extraction phase); second, the entire model is fine-tuned with a lower learning "
    "rate to allow the pre-trained backbone weights to adapt to the specific characteristics "
    "of the target dataset (fine-tuning phase). A learning rate scheduler is employed to "
    "gradually reduce the learning rate as training progresses, improving final model "
    "performance. Early stopping is also implemented, monitoring validation loss to halt "
    "training when the model's performance on unseen data ceases to improve, preventing overfitting."
)

# 3.3.4
add_heading('3.3.4 Evaluation and Persistence', 3)
add_para(
    "Upon completion of training, the model's performance is rigorously evaluated on the "
    "held-out test set using a comprehensive set of metrics. The primary metric is overall "
    "classification accuracy. Additionally, per-class precision, recall, and F1-score are "
    "calculated to provide a nuanced understanding of the model's performance across different "
    "categories, particularly important when class imbalance is present. A confusion matrix "
    "is generated to visually identify which classes are most frequently confused with one "
    "another, informing potential future improvements. Once the model achieves satisfactory "
    "performance, it is serialized and saved to disk using PyTorch's model serialization "
    "functionality (saving the model's state dictionary). The saved model file is then "
    "loaded by the Django back-end at runtime to serve real-time predictions through "
    "the AI Classification API endpoint."
)

# ─────────────────────────────────────────
# 3.4 Dashboard Section
# ─────────────────────────────────────────
add_heading('3.4 Dashboard Section', 2)
add_para(
    "The administrative Dashboard is the central control and analytics hub for system administrators "
    "and journal managers. It provides a real-time, data-rich interface that aggregates key "
    "platform metrics into actionable insights. Built with interactive data visualization "
    "libraries, the dashboard transforms raw database statistics into clear charts, graphs, "
    "and tables that support informed decision-making."
)

# 3.4.1
add_heading('3.4.1 System Components', 3)
add_para(
    "The dashboard is composed of several key components, each designed to surface specific "
    "operational and analytical information:"
)
add_bullet("KPI Summary Cards: A row of at-a-glance metric cards displayed at the top of the dashboard, showing critical totals such as Total Registered Users, Total Articles Submitted, Total Articles Published, and Pending Reviews. These provide an instant system health overview.")
add_bullet("User Activity Charts: Line and bar charts displaying user registration trends over time, daily and monthly active users, and login frequency. These visualizations help administrators understand platform engagement and growth patterns.")
add_bullet("Submission Pipeline Funnel: A funnel chart that visualizes the article submission and review workflow, showing how many articles are at each stage (Submitted → Under Review → Revision → Accepted → Published). This component helps editors manage their workload and identify bottlenecks in the editorial process.")
add_bullet("Sales by Category Chart: An interactive bar or pie chart breaking down platform activity (e.g., article submissions or user interest) by academic category or major, allowing administrators to identify the most popular fields of study on the platform.")
add_bullet("Treemap Visualization: A hierarchical treemap chart that allows administrators to drill down into data organized by Region, Category, and Sub-Category. This component provides a powerful multi-dimensional view of the platform's geographic and academic distribution.")
add_bullet("Content Management Panel: An interface for administrators to manage the platform's master data, including adding, editing, and deactivating journals, managing university and major catalogs, and overseeing user accounts and roles.")
add_bullet("System Logs and Audit Trail: A searchable log of recent system events, including user logins, article submissions, editorial decisions, and administrative actions. This component is critical for security monitoring and system auditing.")

add_para(
    "Together, these components provide a holistic and powerful administrative interface that "
    "gives system managers the visibility and control they need to ensure the platform operates "
    "smoothly, efficiently, and in alignment with its academic mission."
)

doc.save(output_path)
print(f'Document saved to {output_path}')
