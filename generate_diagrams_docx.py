"""
Script to generate a Word document containing all project diagrams
for the University Advisor System project report.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Image paths ────────────────────────────────────────────────────────────
BRAIN_DIR = r"C:\Users\ALAOI\.gemini\antigravity\brain\83e5023f-855a-4fe1-afea-3e4d4f567139"

IMAGES = [
    {
        "path": os.path.join(BRAIN_DIR, "figure_2_1_use_case_frontend_1778135838782.png"),
        "caption": "Figure 2.1: Use Case Diagram – University Advisor System Front-End",
        "description": (
            "This diagram illustrates the functional interactions between system actors "
            "(Student and Guest User) and the University Advisor System's front-end modules. "
            "It captures all primary use cases including authentication, course browsing, "
            "AI chat, GPA tracking, and journal access."
        ),
    },
    {
        "path": os.path.join(BRAIN_DIR, "figure_2_2_flowchart_journal_1778136020306.png"),
        "caption": "Figure 2.2: Flowchart – Scientific Journal Section Workflow",
        "description": (
            "This flowchart details the step-by-step process a user follows when interacting "
            "with the Scientific Journal section, including article search, authentication checks, "
            "bookmarking, and commenting functionality."
        ),
    },
    {
        "path": os.path.join(BRAIN_DIR, "figure_2_4_role_based_diagram_1778136056306.png"),
        "caption": "Figure 2.4: Role-Based System Behavior Diagram",
        "description": (
            "A swimlane diagram depicting the distinct behavioral responsibilities of each "
            "system role (Student, System, Admin). It highlights how requests are routed and "
            "processed based on role-based access control (RBAC) logic."
        ),
    },
    {
        "path": os.path.join(BRAIN_DIR, "figure_2_5_use_case_dashboard_1778136171832.png"),
        "caption": "Figure 2.5: Use Case Diagram – Administrative Dashboard",
        "description": (
            "This diagram maps the administrative use cases available to Super Admin and System Admin "
            "roles within the Administrative Dashboard, covering user management, analytics, "
            "content control, and system configuration."
        ),
    },
    {
        "path": os.path.join(BRAIN_DIR, "figure_2_6_data_pipeline_flowchart_1778136240001.png"),
        "caption": "Figure 2.6: System Workflow Flowchart – Dashboard Data Pipeline",
        "description": (
            "This flowchart illustrates the end-to-end data pipeline powering the Administrative "
            "Dashboard, from raw data collection across system modules, through aggregation and "
            "validation, to final rendering of analytics and exportable reports."
        ),
    },
    {
        "path": os.path.join(BRAIN_DIR, "project_schedule_gantt_1778136296029.png"),
        "caption": "Project Schedule – University Advisor System",
        "description": (
            "A Gantt chart representing the planned project timeline across 16 weeks, covering all "
            "phases from requirements analysis and system design through development, testing, "
            "deployment, and final submission."
        ),
    },
]

OUTPUT_PATH = r"C:\Users\ALAOI\university_advisor\Project_Diagrams.docx"


# ─── Helpers ────────────────────────────────────────────────────────────────

def set_paragraph_border_bottom(paragraph):
    """Add a thin bottom border to a paragraph (used under section titles)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E4A7A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_cover_page(doc: Document):
    """Add a styled cover page."""
    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("University Advisor System")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)

    doc.add_paragraph()  # spacer

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Diagrams & System Models")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x6E, 0xAD)

    doc.add_paragraph()

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 55)
    run.font.color.rgb = RGBColor(0x2E, 0x6E, 0xAD)

    doc.add_paragraph()

    # Description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document contains all system diagrams referenced in Chapter 2\n"
        "of the University Advisor System Final Project Report.\n\n"
        "Diagrams include Use Case Diagrams, Flowcharts,\n"
        "Role-Based Behavior Diagrams, and the Project Schedule."
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Academic Year 2025 – 2026")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_toc(doc: Document):
    """Add a simple Table of Contents."""
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    set_paragraph_border_bottom(p)

    doc.add_paragraph()

    entries = [
        ("Figure 2.1", "Use Case Diagram – University Advisor System Front-End", "3"),
        ("Figure 2.2", "Flowchart – Scientific Journal Section Workflow", "4"),
        ("Figure 2.4", "Role-Based System Behavior Diagram", "5"),
        ("Figure 2.5", "Use Case Diagram – Administrative Dashboard", "6"),
        ("Figure 2.6", "System Workflow Flowchart – Dashboard Data Pipeline", "7"),
        ("", "Project Schedule – University Advisor System", "8"),
    ]

    for fig, title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        # Figure label
        if fig:
            run = p.add_run(f"{fig}: ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2E, 0x6E, 0xAD)
        # Title
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        # Dots + page
        run = p.add_run(f"  {'.' * 60}  {page}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def add_diagram(doc: Document, img: dict):
    """Add a single diagram with caption and description."""
    # Caption / heading
    p = doc.add_paragraph()
    run = p.add_run(img["caption"])
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_border_bottom(p)

    doc.add_paragraph()  # small spacer

    # Image
    img_path = img["path"]
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(6.0))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Image not found: {os.path.basename(img_path)}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        print(f"  ⚠ WARNING: Image not found → {img_path}")

    doc.add_paragraph()  # spacer

    # Description
    p = doc.add_paragraph()
    run = p.add_run(img["description"])
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.italic = True
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    print("Creating Word document with project diagrams...")

    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Default paragraph font ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Build document ────────────────────────────────────────────────────────
    add_cover_page(doc)
    add_toc(doc)

    for img in IMAGES:
        print(f"  Adding: {img['caption']}")
        add_diagram(doc, img)

    # Remove last page break (optional polish)
    # doc.paragraphs[-1].clear()

    doc.save(OUTPUT_PATH)
    print(f"\nDONE! Document saved to:\n   {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
