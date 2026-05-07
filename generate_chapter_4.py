import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_path = r'C:\Users\ALAOI\university_advisor\Chapter_4_English.docx'
doc = docx.Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def b(text):
    doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    para = doc.add_paragraph(text)
    para.style = 'Normal'
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = docx.shared.Pt(9)

# ── CHAPTER 4 ──────────────────────────────────────────
h('Chapter 4', 1)
h('Implementation', 2)
p("This chapter presents the concrete implementation details of the University Advisor System across all four components: Front-End, Back-End, AI, and Dashboard. It documents the actual code structures, design patterns, and technical decisions applied during the development phase, translating the designs from Chapter 3 into a fully functional system.")

# ══════════════════════════════════════════════════════
# 4.1 Front-End Section
# ══════════════════════════════════════════════════════
h('4.1 Front-End Section', 2)
p("The front-end was implemented using Django's template inheritance system, Tailwind CSS utility classes, and vanilla JavaScript. A master base.html template defines the global layout — including the responsive navigation bar, footer, and script/style imports — which all other templates extend. This ensures visual consistency and eliminates code redundancy across the platform.")

p("Key implementation highlights include:")
b("Responsive Navigation Bar: Built with Tailwind's Flexbox utilities and a JavaScript-controlled hamburger menu for mobile views. The nav dynamically renders links based on the authenticated user's role using Django's template tags ({% if user.role == 'Editor' %}).")
b("Article Cards: The journal listing page uses a CSS Grid layout to display article cards. Each card is built with Tailwind utility classes for shadow, border-radius, padding, and hover transitions (hover:scale-105, hover:shadow-xl), creating an interactive and modern feel.")
b("Dynamic Forms: The article submission and login forms use JavaScript for real-time client-side validation. Input fields display inline error messages and visual indicators (red border on invalid, green on valid) without a page reload, using the input and blur event listeners.")
b("Fetch API Integration: The AI classification result page uses the Fetch API to submit the uploaded image to the Django back-end asynchronously and render the prediction results dynamically within the page, providing a smooth single-page-like experience.")
b("Tailwind Configuration: The tailwind.config.js file defines a custom color palette aligned with the University Advisor branding — primary blues and greens for academic professionalism — along with custom font families loaded from Google Fonts (Inter, Outfit).")

p("[Figure 4.1: Screenshot – Homepage with Responsive Navigation]")
p("[Figure 4.2: Screenshot – Scientific Journal Article Listing Page]")
p("[Figure 4.3: Screenshot – Article Submission Form with Validation]")

# ══════════════════════════════════════════════════════
# 4.2 Back-End Section
# ══════════════════════════════════════════════════════
h('4.2 Back-End Section', 2)
p("The back-end was implemented using Django's MVT architecture. Models define the data schema, views contain the business logic, and URL configurations route incoming requests. The following subsections detail the implementation of the core data models.")

# 4.2.1 Custom User Model
h('4.2.1 Custom User Model', 3)
p("A custom user model was implemented by extending Django's AbstractUser class. This approach is a Django best practice that allows the user model to be extended at any point without complex database migrations. The custom model adds a role field to distinguish between the system's four user types.")

code_block(
"""from django.contrib.auth.models import AbstractUser
from django.db import models

class User(AbstractUser):
    ROLE_CHOICES = [
        ('student',       'Student'),
        ('author',        'Author'),
        ('editor',        'Editor'),
        ('administrator', 'Administrator'),
    ]
    role = models.CharField(
        max_length=20,
        choices=ROLE_CHOICES,
        default='student'
    )
    institution = models.CharField(max_length=255, blank=True, null=True)
    bio = models.TextField(blank=True, null=True)

    def is_editor(self):
        return self.role == 'editor'

    def is_author(self):
        return self.role == 'author'

    def is_administrator(self):
        return self.role == 'administrator'

    def __str__(self):
        return f"{self.username} ({self.get_role_display()})"

# In settings.py:
# AUTH_USER_MODEL = 'accounts.User'"""
)

p("The AUTH_USER_MODEL setting in settings.py is pointed to this custom model before any migrations are run, which is a critical requirement when using a custom user model in Django. The helper methods (is_editor, is_author, is_administrator) are used throughout views and templates to conditionally render role-specific content.")

# 4.2.2 Author Model
h('4.2.2 Author Model', 3)
p("The Author Model extends the custom User model through a OneToOneField relationship. This profile-based extension pattern keeps the core User model lean while allowing author-specific attributes to be added cleanly. The model stores academic metadata relevant to a researcher's profile.")

code_block(
"""from django.db import models
from accounts.models import User

class Author(models.Model):
    user = models.OneToOneField(
        User,
        on_delete=models.CASCADE,
        related_name='author_profile'
    )
    orcid_id = models.CharField(
        max_length=19,
        blank=True,
        null=True,
        help_text="Format: 0000-0000-0000-0000"
    )
    research_interests = models.TextField(blank=True, null=True)
    h_index = models.PositiveIntegerField(default=0)
    total_publications = models.PositiveIntegerField(default=0)

    def __str__(self):
        return f"Author: {self.user.get_full_name()}"

    def get_published_articles(self):
        return self.articles.filter(status='published')

    def get_pending_articles(self):
        return self.articles.filter(
            status__in=['submitted', 'under_review', 'revision_requested']
        )

    class Meta:
        ordering = ['user__last_name']"""
)

p("A Django signal (post_save) is used to automatically create an Author profile instance whenever a new User account with the role of 'author' is registered. This ensures that the one-to-one relationship is always maintained without requiring manual profile creation. The get_published_articles and get_pending_articles helper methods encapsulate common queries, promoting code reuse across views and templates.")

p("[Figure 4.4: Django Admin View – Custom User Model Management]")
p("[Figure 4.5: Django Admin View – Author Profile Management]")

# ══════════════════════════════════════════════════════
# 4.3 AI Section
# ══════════════════════════════════════════════════════
h('4.3 AI Section', 2)
p("The AI image classification module was implemented in Python using the PyTorch deep learning framework. The implementation follows a structured pipeline from data ingestion through to model deployment, as described in the following subsections.")

# 4.3.1
h('4.3.1 Imports and Initial Setup', 3)
code_block(
"""import os
import copy
import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils.data import Dataset, DataLoader
from torchvision import models, transforms
from torchvision.models import ResNet50_Weights
from PIL import Image
from sklearn.model_selection import train_test_split
from sklearn.metrics import classification_report, confusion_matrix
import numpy as np
import matplotlib.pyplot as plt

# ── Device Configuration ──────────────────────────────
DEVICE = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
print(f'Training on: {DEVICE}')

# ── Hyperparameters ───────────────────────────────────
IMG_SIZE    = 224
BATCH_SIZE  = 32
HEAD_EPOCHS = 10
FINE_EPOCHS = 20
HEAD_LR     = 1e-3
FINE_LR     = 1e-4
NUM_CLASSES = 5   # number of target categories
DATA_DIR    = 'dataset/'
MODEL_PATH  = 'classifier_model.pth'"""
)
p("The setup block detects the available hardware and configures the device accordingly — using CUDA-enabled GPU if available, falling back to CPU otherwise. All hyperparameters are defined as constants at the top of the file to make experimentation and tuning straightforward.")

# 4.3.2
h('4.3.2 Dataset Construction and Preprocessing', 3)
code_block(
"""# ── Collect file paths and labels ────────────────────
all_paths, all_labels = [], []
class_names = sorted(os.listdir(DATA_DIR))
class_to_idx = {cls: idx for idx, cls in enumerate(class_names)}

for cls in class_names:
    cls_dir = os.path.join(DATA_DIR, cls)
    for fname in os.listdir(cls_dir):
        if fname.lower().endswith(('.jpg', '.jpeg', '.png')):
            all_paths.append(os.path.join(cls_dir, fname))
            all_labels.append(class_to_idx[cls])

# ── Stratified train / validation / test split ────────
train_p, temp_p, train_l, temp_l = train_test_split(
    all_paths, all_labels, test_size=0.3,
    stratify=all_labels, random_state=42
)
val_p, test_p, val_l, test_l = train_test_split(
    temp_p, temp_l, test_size=0.5,
    stratify=temp_l, random_state=42
)
print(f"Train: {len(train_p)} | Val: {len(val_p)} | Test: {len(test_p)}")"""
)
p("The dataset directory is expected to follow the standard ImageFolder structure, where each subdirectory represents a class. The stratified split ensures that the class distribution is proportionally maintained across all three subsets, which is critical for unbiased model evaluation, especially when class imbalance is present.")

# 4.3.3
h('4.3.3 Data Augmentation and DataLoader', 3)
code_block(
"""# ── Transform pipelines ───────────────────────────────
train_tf = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.RandomHorizontalFlip(p=0.5),
    transforms.RandomVerticalFlip(p=0.2),
    transforms.ColorJitter(brightness=0.3, contrast=0.3,
                           saturation=0.2, hue=0.1),
    transforms.RandomRotation(degrees=15),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225]),
])
val_test_tf = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225]),
])"""
)
p("Data augmentation is applied only to the training set. Random flips, color jitter, and rotations artificially increase the diversity of training examples, reducing overfitting. Normalization values (mean and std) match those used during ImageNet pre-training, ensuring compatibility with the pre-trained ResNet50 backbone. The validation and test pipelines apply only resizing and normalization to ensure consistent evaluation.")

h('4.3.3.1 Initialization Method', 4)
code_block(
"""class ImageDataset(Dataset):
    def __init__(self, paths, labels, transform=None):
        self.paths     = paths
        self.labels    = labels
        self.transform = transform"""
)
p("The __init__ method stores the list of image file paths, their corresponding integer class labels, and the transform pipeline. This constructor is called once when the dataset object is instantiated, keeping the initialization lightweight.")

h('4.3.3.2 Item Retrieval Method', 4)
code_block(
"""    def __getitem__(self, idx):
        img = Image.open(self.paths[idx]).convert('RGB')
        if self.transform:
            img = self.transform(img)
        return img, self.labels[idx]"""
)
p("The __getitem__ method is called by the DataLoader for each sample. It loads the image from disk using PIL, converts it to RGB (handling grayscale or RGBA images), applies the transform pipeline, and returns the processed tensor alongside its integer label.")

h('4.3.3.3 Length Method', 4)
code_block(
"""    def __len__(self):
        return len(self.paths)

# ── Instantiate Datasets and DataLoaders ─────────────
train_ds = ImageDataset(train_p, train_l, train_tf)
val_ds   = ImageDataset(val_p,   val_l,   val_test_tf)
test_ds  = ImageDataset(test_p,  test_l,  val_test_tf)

train_dl = DataLoader(train_ds, batch_size=BATCH_SIZE,
                      shuffle=True,  num_workers=2, pin_memory=True)
val_dl   = DataLoader(val_ds,   batch_size=BATCH_SIZE,
                      shuffle=False, num_workers=2, pin_memory=True)
test_dl  = DataLoader(test_ds,  batch_size=BATCH_SIZE,
                      shuffle=False, num_workers=2)"""
)
p("The __len__ method returns the total number of samples. The DataLoaders wrap the datasets and handle batching, shuffling (training only), and parallel data loading via multiple worker processes. pin_memory=True is set for GPU training to accelerate the CPU-to-GPU memory transfer.")

# 4.3.4
h('4.3.4 Model Initialization and Transfer Learning', 3)
code_block(
"""# ── Load pre-trained ResNet50 ─────────────────────────
backbone = models.resnet50(weights=ResNet50_Weights.IMAGENET1K_V2)

# Freeze all backbone parameters
for param in backbone.parameters():
    param.requires_grad = False

# ── Replace the final classification head ─────────────
in_features = backbone.fc.in_features
backbone.fc = nn.Sequential(
    nn.Linear(in_features, 512),
    nn.ReLU(),
    nn.Dropout(p=0.4),
    nn.Linear(512, NUM_CLASSES)
)

model = backbone.to(DEVICE)

# ── Loss and Optimizer (Head-only phase) ──────────────
criterion = nn.CrossEntropyLoss()
optimizer = optim.Adam(model.fc.parameters(), lr=HEAD_LR)
scheduler = optim.lr_scheduler.StepLR(
    optimizer, step_size=5, gamma=0.5
)"""
)
p("ResNet50 pre-trained on ImageNet is loaded and all its convolutional parameters are frozen, preserving the powerful feature representations learned during pre-training. Only the newly added classification head — a two-layer MLP with ReLU activation, Dropout for regularization, and a final linear layer sized to NUM_CLASSES — is trainable during the initial head-training phase. The Adam optimizer and a step-based learning rate scheduler are configured.")

# 4.3.5
h('4.3.5 Training Loop', 3)
h('4.3.5.1 Head Training Phase', 4)
code_block(
"""def train_one_epoch(model, loader, optimizer, criterion):
    model.train()
    total_loss, correct = 0.0, 0
    for imgs, labels in loader:
        imgs, labels = imgs.to(DEVICE), labels.to(DEVICE)
        optimizer.zero_grad()
        outputs = model(imgs)
        loss = criterion(outputs, labels)
        loss.backward()
        optimizer.step()
        total_loss += loss.item() * imgs.size(0)
        correct += (outputs.argmax(1) == labels).sum().item()
    return total_loss / len(loader.dataset), correct / len(loader.dataset)

print("── Phase 1: Training Classification Head ──")
for epoch in range(HEAD_EPOCHS):
    tr_loss, tr_acc = train_one_epoch(model, train_dl, optimizer, criterion)
    scheduler.step()
    print(f"Epoch {epoch+1}/{HEAD_EPOCHS} | Loss: {tr_loss:.4f} | Acc: {tr_acc:.4f}")"""
)
p("The head-training phase trains only the newly added classification head while the backbone remains frozen. This phase quickly converges the head weights to produce a reasonable initial classification boundary, which is essential before unfreezing the backbone for fine-tuning.")

h('4.3.5.2 Fine-Tuning Phase', 4)
code_block(
"""# ── Unfreeze all backbone layers for fine-tuning ─────
for param in model.parameters():
    param.requires_grad = True

fine_optimizer = optim.Adam(model.parameters(), lr=FINE_LR)
fine_scheduler = optim.lr_scheduler.CosineAnnealingLR(
    fine_optimizer, T_max=FINE_EPOCHS
)

best_val_acc  = 0.0
best_model_wts = copy.deepcopy(model.state_dict())

print("── Phase 2: Fine-Tuning Full Network ──")
for epoch in range(FINE_EPOCHS):
    tr_loss, tr_acc = train_one_epoch(
        model, train_dl, fine_optimizer, criterion
    )
    val_loss, val_acc = evaluate(model, val_dl, criterion)
    fine_scheduler.step()

    if val_acc > best_val_acc:
        best_val_acc  = val_acc
        best_model_wts = copy.deepcopy(model.state_dict())

    print(f"Epoch {epoch+1}/{FINE_EPOCHS} | "
          f"Train Acc: {tr_acc:.4f} | Val Acc: {val_acc:.4f}")

print(f"Best Validation Accuracy: {best_val_acc:.4f}")
model.load_state_dict(best_model_wts)"""
)
p("During fine-tuning, all layers are unfrozen and trained with a much lower learning rate to gently adapt the pre-trained backbone features to the target domain without catastrophic forgetting. A Cosine Annealing scheduler smoothly decays the learning rate. Best model weights are tracked using the validation accuracy, and the best checkpoint is restored at the end of training.")

# 4.3.6
h('4.3.6 Evaluation, Saving, and Inference', 3)
code_block(
"""def evaluate(model, loader, criterion):
    model.eval()
    total_loss, correct = 0.0, 0
    all_preds, all_labels = [], []
    with torch.no_grad():
        for imgs, labels in loader:
            imgs, labels = imgs.to(DEVICE), labels.to(DEVICE)
            outputs = model(imgs)
            loss = criterion(outputs, labels)
            total_loss += loss.item() * imgs.size(0)
            preds = outputs.argmax(1)
            correct += (preds == labels).sum().item()
            all_preds.extend(preds.cpu().numpy())
            all_labels.extend(labels.cpu().numpy())
    return total_loss / len(loader.dataset), correct / len(loader.dataset)

# ── Final Test Evaluation ─────────────────────────────
_, test_acc = evaluate(model, test_dl, criterion)
print(f"Test Accuracy: {test_acc:.4f}")

# ── Classification Report ─────────────────────────────
# (run evaluate with preds collection for full report)
# print(classification_report(all_labels, all_preds,
#                             target_names=class_names))

# ── Save Model ────────────────────────────────────────
torch.save(model.state_dict(), MODEL_PATH)
print(f"Model saved to {MODEL_PATH}")

# ── Inference Function ────────────────────────────────
def predict_image(image_path, model, transform, class_names):
    model.eval()
    img = Image.open(image_path).convert('RGB')
    tensor = transform(img).unsqueeze(0).to(DEVICE)
    with torch.no_grad():
        output = model(tensor)
        probs  = torch.softmax(output, dim=1).squeeze()
        pred   = probs.argmax().item()
    return class_names[pred], probs[pred].item(), probs.cpu().numpy()"""
)
p("The evaluate function runs the model in inference mode (model.eval()) and disables gradient computation (torch.no_grad()) for efficiency. After final test evaluation, the best model's state dictionary is serialized to disk. The predict_image function provides a clean interface for single-image inference, returning the predicted class name, confidence score, and the full probability distribution across all classes — which is used by the Django API endpoint to populate the classification result page.")

# ══════════════════════════════════════════════════════
# 4.4 Dashboard Section
# ══════════════════════════════════════════════════════
h('4.4 Dashboard Section', 2)
p("The administrative dashboard was implemented using Streamlit, with Pandas for data processing and Plotly for interactive visualizations. The dashboard connects directly to the PostgreSQL database using SQLAlchemy and renders real-time analytics for administrators.")

p("The main dashboard application (dashboard.py) is structured as follows:")

code_block(
"""import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from sqlalchemy import create_engine
import openpyxl
from io import BytesIO

# ── Page Configuration ────────────────────────────────
st.set_page_config(
    page_title="University Advisor – Admin Dashboard",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── Database Connection ───────────────────────────────
@st.cache_resource
def get_engine():
    DB_URL = "postgresql://user:password@localhost:5432/university_advisor"
    return create_engine(DB_URL)

engine = get_engine()

# ── Data Loading ──────────────────────────────────────
@st.cache_data(ttl=300)
def load_data():
    articles_df = pd.read_sql("SELECT * FROM journal_article", engine)
    users_df    = pd.read_sql("SELECT * FROM accounts_user",   engine)
    reviews_df  = pd.read_sql("SELECT * FROM journal_review",  engine)
    return articles_df, users_df, reviews_df

articles_df, users_df, reviews_df = load_data()"""
)

p("The @st.cache_resource decorator ensures the database engine is created only once and reused across all user sessions. The @st.cache_data(ttl=300) decorator caches the loaded DataFrames for 5 minutes (300 seconds), preventing redundant database queries on every page interaction while still reflecting reasonably current data.")

code_block(
"""# ── Sidebar Filters ───────────────────────────────────
st.sidebar.title("🔍 Filters")
date_range = st.sidebar.date_input("Date Range", [])
selected_status = st.sidebar.multiselect(
    "Article Status",
    options=articles_df['status'].unique(),
    default=articles_df['status'].unique()
)
filtered_df = articles_df[articles_df['status'].isin(selected_status)]

# ── KPI Summary Row ───────────────────────────────────
col1, col2, col3, col4 = st.columns(4)
col1.metric("👥 Total Users",      len(users_df))
col2.metric("📄 Total Articles",   len(articles_df))
col3.metric("✅ Published",        len(articles_df[articles_df['status']=='published']))
col4.metric("⏳ Pending Review",   len(articles_df[articles_df['status']=='under_review']))

# ── Charts ────────────────────────────────────────────
st.subheader("📊 Articles by Status")
status_counts = filtered_df['status'].value_counts().reset_index()
fig_bar = px.bar(status_counts, x='status', y='count',
                 color='status', title='Article Distribution by Status')
st.plotly_chart(fig_bar, use_container_width=True)

st.subheader("🗺️ Submissions by Region and Category")
fig_tree = px.treemap(filtered_df, path=['region', 'category', 'sub_category'],
                      values='id', title='Treemap: Region > Category > Sub-Category')
st.plotly_chart(fig_tree, use_container_width=True)

# ── Excel Export ──────────────────────────────────────
def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Articles')
    return output.getvalue()

st.download_button(
    label="📥 Export to Excel",
    data=to_excel(filtered_df),
    file_name="university_advisor_report.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)"""
)

p("The sidebar provides dynamic filter controls that instantly update all charts through Streamlit's reactive execution model. The KPI row uses st.metric components to display key totals with automatic delta indicators. The Plotly treemap provides a hierarchical drill-down visualization from Region to Category to Sub-Category. The Excel export button uses Openpyxl to generate a formatted .xlsx file in memory (using BytesIO) and delivers it directly to the administrator's browser as a downloadable file, requiring no server-side file storage.")

p("[Figure 4.6: Screenshot – Administrative Dashboard KPI Summary and Article Status Chart]")
p("[Figure 4.7: Screenshot – Dashboard Treemap Visualization (Region > Category > Sub-Category)]")
p("[Figure 4.8: Screenshot – Dashboard Excel Export Functionality]")

doc.save(output_path)
print(f'Document saved to {output_path}')
