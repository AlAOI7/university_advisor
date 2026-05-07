import docx
from docx.shared import Pt, Inches
import os
import re

doc_path = r'C:\Users\ALAOI\university_advisor\Muthnab ScienceNet Graduation Project.2 (1) (1).docx'
output_path = r'C:\Users\ALAOI\university_advisor\University_Advisor_Final_English.docx'

doc = docx.Document(doc_path)

# Update Title Page
doc.paragraphs[16].text = 'University Advisor System'
doc.paragraphs[19].text = 'Juri Mutlaq Al-Mutairi\nFajr Al-Humaidi Al-Mutairi\nFouz Faisal Al-Otaibi\nDana Fahad Al-Suwaidani\nMunira Sulaiman Al-Sadrani\nHissa Mohammed Al-Harabi'
doc.paragraphs[22].text = 'Supervisor: Wedad Al-Awad'

# Let's clear out all paragraphs after the table of contents (approx paragraph 50)
# This is a bit aggressive, but we need to remove the old project's data.
# We will keep the first 30 paragraphs (Cover page + TOC header)
for i in range(len(doc.paragraphs)-1, 30, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# Now let's add the content in English
def add_heading(text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')

def add_para(text):
    p = doc.add_paragraph(text, style='Normal')

# Abstract
add_heading('Abstract', 1)
add_para("This graduation project presents the 'University Advisor System', an innovative, AI-driven web application designed to guide high school graduates in selecting the most suitable university major. Choosing a major is a critical decision that impacts a student's future career and personal satisfaction. To facilitate this process, our platform integrates a comprehensive aptitude and personality assessment, mapping user traits to potential academic paths. The system features a robust database of academic majors, providing detailed information on required courses, recommended reading materials, and career prospects. To deliver a seamless user experience, the front-end was developed using modern web technologies including HTML5, CSS3, JavaScript, and Tailwind CSS. The robust back-end infrastructure is powered by the Django framework and PostgreSQL, ensuring secure data management and high performance. Furthermore, the integration of advanced Artificial Intelligence models enhances the recommendation engine, providing personalized academic advising. A comprehensive dashboard provides administrators with real-time analytics and management tools. Ultimately, this project demonstrates the effective application of modern software engineering principles to solve real-world educational challenges.")

# Chapter 1: Introduction
add_heading('Chapter 1', 1)
add_heading('Introduction', 2)
add_para("The transition from high school to university is a pivotal moment in a student's educational journey. The sheer variety of academic disciplines, coupled with a lack of personalized guidance, often leaves students feeling overwhelmed and uncertain about their future. The 'University Advisor System' was conceptualized to bridge this gap by providing an intelligent, accessible, and user-centric platform that empowers students to make informed academic choices. By leveraging cutting-edge web development frameworks and artificial intelligence, the system offers a holistic advising experience. This chapter provides an in-depth overview of the core components that constitute the system's architecture, divided into the Front-End, Back-End, AI, and Dashboard sections.")

add_heading('1.1 Front-End Section', 2)
add_para("The Front-End section of the University Advisor System is designed with a primary focus on user experience (UX) and user interface (UI) aesthetics. Recognizing that the target audience primarily consists of young adults, the interface is modern, intuitive, and highly responsive across various devices, including desktop computers, tablets, and smartphones. We utilized HTML5 for semantic structure, CSS3 for advanced styling, and JavaScript for dynamic interactivity. To ensure a consistent and rapid development process, modern utility-first CSS frameworks and robust UI libraries were employed, allowing for highly customizable and visually appealing components. The front-end communicates seamlessly with the back-end APIs to dynamically render interactive assessments, display tailored major recommendations, and present detailed academic resources in an engaging manner.")

add_heading('1.2 Back-End Section', 2)
add_para("The Back-End section serves as the robust foundation of the University Advisor System, responsible for business logic, data processing, and secure communication with the database. Built upon the Python-based Django framework, the back-end follows the Model-View-Template (MVT) architectural pattern, ensuring a clean separation of concerns. Django's built-in authentication system is utilized to securely manage user accounts, while custom logic handles the scoring of aptitude tests and the generation of recommendations. PostgreSQL was selected as the primary relational database management system due to its reliability, scalability, and support for complex queries. The back-end also exposes a set of RESTful APIs, facilitating seamless data exchange with the front-end and paving the way for future mobile application integration. Strict security measures, including CSRF protection and data validation, are implemented to safeguard user information.")

add_heading('1.3 AI Section', 2)
add_para("The Artificial Intelligence (AI) section is a distinguishing feature of the University Advisor System, elevating it from a standard directory to an intelligent advising platform. The core of the AI functionality revolves around a sophisticated recommendation engine that analyzes the user's responses to the aptitude and personality tests. By employing machine learning algorithms, potentially utilizing frameworks like PyTorch or integrating with advanced language models, the system can identify complex patterns in user data and accurately match them with appropriate academic majors. Furthermore, the AI section powers an interactive chatbot, acting as a virtual academic advisor capable of answering user queries in real-time, explaining the rationale behind specific recommendations, and providing general guidance on university life. This personalized AI-driven approach significantly enhances the value provided to the students.")

add_heading('1.4 Dashboard Section', 2)
add_para("The Dashboard section provides a comprehensive administrative interface for system managers and educational counselors. Built with a focus on data visualization and ease of use, the dashboard offers real-time insights into system usage, user demographics, and test completion rates. Administrators can efficiently manage the catalog of universities, academic majors, and associated courses through an intuitive Content Management System (CMS) interface. The dashboard also features analytical tools that aggregate user data to identify trending majors and areas where students may require additional resources. This centralized control panel empowers administrators to maintain the accuracy of the platform's information, monitor overall system health, and make data-driven decisions to continuously improve the advising experience.")

doc.save(output_path)
print(f'Document saved to {output_path}')
