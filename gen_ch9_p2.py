import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

out = r'C:\Users\ALAOI\university_advisor\Chapter_9_English.docx'

# Load the existing doc (Part 1 already saved it)
doc = docx.Document(out)

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

# ═══════════════════════════════════
h('9.3 AI Section', 2)
p("The following listings present the complete deep learning pipeline used to train and deploy the image classification model.")

h('A.3.1 – Library Imports for Deep Learning Pipeline', 3)
code("""import os, copy, random
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image
import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils.data import Dataset, DataLoader
from torchvision import models, transforms
from torchvision.models import ResNet18_Weights
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import (classification_report,
                              confusion_matrix, accuracy_score)

# Reproducibility
SEED = 42
random.seed(SEED); np.random.seed(SEED); torch.manual_seed(SEED)
if torch.cuda.is_available():
    torch.cuda.manual_seed_all(SEED)

DEVICE = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
print(f"Device: {DEVICE}")""")

h('A.3.2 – Image and Label Extraction from Directory Structure', 3)
code("""DATA_DIR = 'dataset/'   # Each sub-folder = one class

image_paths, labels = [], []
class_names = sorted(os.listdir(DATA_DIR))

for class_name in class_names:
    class_folder = os.path.join(DATA_DIR, class_name)
    if not os.path.isdir(class_folder):
        continue
    for fname in os.listdir(class_folder):
        if fname.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp')):
            image_paths.append(os.path.join(class_folder, fname))
            labels.append(class_name)

print(f"Total images found : {len(image_paths)}")
print(f"Classes detected   : {class_names}")""")

h('A.3.3 – Creating a DataFrame for Image Paths and Labels', 3)
code("""df = pd.DataFrame({'image_path': image_paths, 'label': labels})

# Class distribution
print("\\nClass distribution:")
print(df['label'].value_counts())

# Quick sanity check – verify all files exist
missing = [p for p in df['image_path'] if not os.path.isfile(p)]
if missing:
    print(f"WARNING: {len(missing)} missing files detected!")
else:
    print("All image files verified successfully.")

print(df.head())""")

h('A.3.4 – Splitting Dataset and Label Encoding', 3)
code("""le = LabelEncoder()
df['label_encoded'] = le.fit_transform(df['label'])
NUM_CLASSES = len(le.classes_)
print(f"Number of classes: {NUM_CLASSES}")
print(f"Class mapping: {dict(zip(le.classes_, le.transform(le.classes_)))}")

# Stratified 70 / 15 / 15 split
train_df, temp_df = train_test_split(
    df, test_size=0.30, stratify=df['label_encoded'], random_state=SEED)
val_df, test_df = train_test_split(
    temp_df, test_size=0.50, stratify=temp_df['label_encoded'], random_state=SEED)

print(f"Train : {len(train_df):>5} images")
print(f"Val   : {len(val_df):>5} images")
print(f"Test  : {len(test_df):>5} images")""")

h('A.3.5 – Data Augmentation Pipeline for Training Images', 3)
code("""IMG_SIZE = 224

train_transform = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.RandomHorizontalFlip(p=0.5),
    transforms.RandomVerticalFlip(p=0.2),
    transforms.RandomRotation(degrees=15),
    transforms.ColorJitter(brightness=0.3, contrast=0.3,
                           saturation=0.2, hue=0.1),
    transforms.RandomGrayscale(p=0.05),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225]),
])

val_test_transform = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225]),
])""")

h('A.3.6 – Custom Dataset Class for Image Loading and Labeling', 3)
code("""class ImageDataset(Dataset):
    def __init__(self, dataframe, transform=None):
        self.paths     = dataframe['image_path'].tolist()
        self.labels    = dataframe['label_encoded'].tolist()
        self.transform = transform

    def __len__(self):
        return len(self.paths)

    def __getitem__(self, idx):
        img   = Image.open(self.paths[idx]).convert('RGB')
        label = self.labels[idx]
        if self.transform:
            img = self.transform(img)
        return img, label

train_ds = ImageDataset(train_df, train_transform)
val_ds   = ImageDataset(val_df,   val_test_transform)
test_ds  = ImageDataset(test_df,  val_test_transform)""")

h('A.3.7 – Hyperparameter Configuration for Training and Fine-Tuning', 3)
code("""# ── Hyperparameters ───────────────────────────────────
BATCH_SIZE  = 32
HEAD_LR     = 1e-3   # Learning rate for head-only phase
FINE_LR     = 1e-4   # Learning rate for full fine-tuning phase
HEAD_EPOCHS = 10     # Epochs to train classification head only
FINE_EPOCHS = 20     # Epochs for full model fine-tuning
WEIGHT_DECAY = 1e-4  # L2 regularisation coefficient
DROPOUT_P    = 0.4   # Dropout probability in classification head
PATIENCE     = 5     # Early stopping patience (fine-tune phase)

print("Hyperparameter Configuration:")
print(f"  Batch Size     : {BATCH_SIZE}")
print(f"  Head LR        : {HEAD_LR}")
print(f"  Fine-Tune LR   : {FINE_LR}")
print(f"  Head Epochs    : {HEAD_EPOCHS}")
print(f"  Fine-Tune Epochs: {FINE_EPOCHS}")
print(f"  Dropout        : {DROPOUT_P}")
print(f"  Early Stop     : {PATIENCE} epochs")""")

h('A.3.8 – DataLoader for Efficient Batch Processing', 3)
code("""train_loader = DataLoader(
    train_ds,
    batch_size  = BATCH_SIZE,
    shuffle     = True,
    num_workers = 2,
    pin_memory  = True
)
val_loader = DataLoader(
    val_ds,
    batch_size  = BATCH_SIZE,
    shuffle     = False,
    num_workers = 2,
    pin_memory  = True
)
test_loader = DataLoader(
    test_ds,
    batch_size  = BATCH_SIZE,
    shuffle     = False,
    num_workers = 2
)

print(f"Batches per epoch – Train: {len(train_loader)}, "
      f"Val: {len(val_loader)}, Test: {len(test_loader)}")""")

h('A.3.9 – Loading a Pre-trained ResNet-18 Model', 3)
code("""# Load ResNet-18 with ImageNet weights
backbone = models.resnet18(weights=ResNet18_Weights.IMAGENET1K_V1)

# Inspect the original classifier head
print("Original FC layer:", backbone.fc)
# Output: Linear(in_features=512, out_features=1000, bias=True)

in_features = backbone.fc.in_features
print(f"Feature vector size: {in_features}")""")

h('A.3.10 – Freezing Backbone and Customizing Output Layer', 3)
code("""# Freeze all backbone parameters
for name, param in backbone.named_parameters():
    param.requires_grad = False

# Replace the classification head with a custom MLP
backbone.fc = nn.Sequential(
    nn.Linear(in_features, 256),
    nn.BatchNorm1d(256),
    nn.ReLU(inplace=True),
    nn.Dropout(p=DROPOUT_P),
    nn.Linear(256, NUM_CLASSES)
)

model = backbone.to(DEVICE)

# Verify: only head params are trainable
trainable = sum(p.numel() for p in model.parameters() if p.requires_grad)
total     = sum(p.numel() for p in model.parameters())
print(f"Trainable params : {trainable:,} / {total:,} "
      f"({100*trainable/total:.1f}%)")""")

h('A.3.11 – Optimizer and Learning Rate Scheduler Setup', 3)
code("""criterion = nn.CrossEntropyLoss(label_smoothing=0.1)

# Phase 1 – head only
head_optimizer = optim.Adam(
    model.fc.parameters(), lr=HEAD_LR, weight_decay=WEIGHT_DECAY)
head_scheduler = optim.lr_scheduler.StepLR(
    head_optimizer, step_size=4, gamma=0.5)

# Phase 2 – full model (initialised after unfreezing)
# fine_optimizer and fine_scheduler defined in A.3.13""")

h('A.3.12 – Training Loop for Classification Head', 3)
code("""def run_epoch(model, loader, optimizer, criterion, training=True):
    model.train() if training else model.eval()
    total_loss, correct, n = 0.0, 0, 0
    with torch.set_grad_enabled(training):
        for imgs, labels in loader:
            imgs, labels = imgs.to(DEVICE), labels.to(DEVICE)
            if training:
                optimizer.zero_grad()
            outputs = model(imgs)
            loss    = criterion(outputs, labels)
            if training:
                loss.backward()
                optimizer.step()
            total_loss += loss.item() * imgs.size(0)
            correct    += (outputs.argmax(1) == labels).sum().item()
            n          += imgs.size(0)
    return total_loss / n, correct / n

print("── Phase 1: Training Classification Head ──")
for epoch in range(1, HEAD_EPOCHS + 1):
    tr_loss, tr_acc = run_epoch(model, train_loader, head_optimizer, criterion)
    vl_loss, vl_acc = run_epoch(model, val_loader,   head_optimizer, criterion,
                                training=False)
    head_scheduler.step()
    print(f"Ep {epoch:02d}/{HEAD_EPOCHS} | "
          f"Train {tr_acc:.4f} | Val {vl_acc:.4f}")""")

h('A.3.13 – Enabling Fine-Tuning of the Full Model', 3)
code("""# Unfreeze all layers
for param in model.parameters():
    param.requires_grad = True

fine_optimizer = optim.AdamW(
    model.parameters(), lr=FINE_LR, weight_decay=WEIGHT_DECAY)
fine_scheduler = optim.lr_scheduler.CosineAnnealingLR(
    fine_optimizer, T_max=FINE_EPOCHS, eta_min=1e-6)

best_val_acc  = 0.0
best_wts      = copy.deepcopy(model.state_dict())
patience_ctr  = 0

print("── Phase 2: Fine-Tuning Full Network ──")
for epoch in range(1, FINE_EPOCHS + 1):
    tr_loss, tr_acc = run_epoch(model, train_loader, fine_optimizer, criterion)
    vl_loss, vl_acc = run_epoch(model, val_loader,   fine_optimizer, criterion,
                                training=False)
    fine_scheduler.step()

    if vl_acc > best_val_acc:
        best_val_acc = vl_acc
        best_wts     = copy.deepcopy(model.state_dict())
        patience_ctr = 0
    else:
        patience_ctr += 1

    print(f"Ep {epoch:02d}/{FINE_EPOCHS} | "
          f"Train {tr_acc:.4f} | Val {vl_acc:.4f}"
          + (" ← best" if patience_ctr == 0 else ""))

    if patience_ctr >= PATIENCE:
        print("Early stopping triggered.")
        break

model.load_state_dict(best_wts)
print(f"Best Validation Accuracy: {best_val_acc:.4f}")""")

h('A.3.14 – Model Evaluation on the Test Set', 3)
code("""model.eval()
all_preds, all_labels_list = [], []

with torch.no_grad():
    for imgs, labels in test_loader:
        imgs   = imgs.to(DEVICE)
        preds  = model(imgs).argmax(1).cpu().numpy()
        all_preds.extend(preds)
        all_labels_list.extend(labels.numpy())

test_acc = accuracy_score(all_labels_list, all_preds)
print(f"\\nTest Accuracy: {test_acc:.4f} ({test_acc*100:.2f}%)")
print("\\nClassification Report:")
print(classification_report(all_labels_list, all_preds,
                             target_names=le.classes_))

# Confusion Matrix
cm = confusion_matrix(all_labels_list, all_preds)
plt.figure(figsize=(8, 6))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=le.classes_, yticklabels=le.classes_)
plt.title('Confusion Matrix – Test Set')
plt.xlabel('Predicted Label')
plt.ylabel('True Label')
plt.tight_layout()
plt.savefig('confusion_matrix.png', dpi=150)
plt.show()""")

h('A.3.15 – Saving Trained Model Weights', 3)
code("""MODEL_PATH = 'classifier_resnet18.pth'

# Save state dict (recommended – framework-agnostic)
torch.save({
    'model_state_dict' : model.state_dict(),
    'class_names'      : list(le.classes_),
    'num_classes'      : NUM_CLASSES,
    'img_size'         : IMG_SIZE,
    'test_accuracy'    : test_acc,
}, MODEL_PATH)

print(f"Model checkpoint saved → {MODEL_PATH}")""")

h('A.3.16 – Loading and Preparing the Trained Model for Inference', 3)
code("""def load_model(checkpoint_path, device):
    ckpt        = torch.load(checkpoint_path, map_location=device)
    num_classes = ckpt['num_classes']

    base = models.resnet18(weights=None)
    base.fc = nn.Sequential(
        nn.Linear(base.fc.in_features, 256),
        nn.BatchNorm1d(256),
        nn.ReLU(inplace=True),
        nn.Dropout(p=0.4),
        nn.Linear(256, num_classes)
    )
    base.load_state_dict(ckpt['model_state_dict'])
    base.eval()
    return base.to(device), ckpt['class_names']

def predict(image_path, model, class_names, device, img_size=224):
    tf = transforms.Compose([
        transforms.Resize((img_size, img_size)),
        transforms.ToTensor(),
        transforms.Normalize([0.485,0.456,0.406],[0.229,0.224,0.225]),
    ])
    img    = Image.open(image_path).convert('RGB')
    tensor = tf(img).unsqueeze(0).to(device)
    with torch.no_grad():
        probs = torch.softmax(model(tensor), dim=1).squeeze()
    pred_idx = probs.argmax().item()
    return {
        'class'      : class_names[pred_idx],
        'confidence' : round(probs[pred_idx].item() * 100, 2),
        'all_probs'  : {c: round(probs[i].item()*100,2)
                        for i, c in enumerate(class_names)},
    }

# Usage:
# model, classes = load_model('classifier_resnet18.pth', DEVICE)
# result = predict('sample.jpg', model, classes, DEVICE)
# print(result)""")

# ═══════════════════════════════════
h('9.4 Dashboard Section', 2)
p("The following listings present the Streamlit dashboard source code for the University Advisor System's administrative analytics panel.")

h('A.4.1 – RTL Interface Configuration for Arabic Support', 3)
code("""import streamlit as st

# ── Page config ───────────────────────────────────────
st.set_page_config(
    page_title="University Advisor – لوحة التحكم",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── RTL CSS injection ─────────────────────────────────
RTL_CSS = \"\"\"
<style>
  /* Force RTL direction on the entire Streamlit app */
  .stApp, [data-testid="stSidebar"],
  [data-testid="stMarkdownContainer"] p,
  [data-testid="stDataFrame"] {
      direction: rtl;
      text-align: right;
      font-family: 'Tajawal', 'Cairo', sans-serif;
  }
  /* Keep charts LTR to avoid axis mirroring */
  .js-plotly-plot, iframe { direction: ltr !important; }
  /* Sidebar right-alignment */
  [data-testid="stSidebar"] .block-container { padding-right: 1rem; }
</style>
<link href="https://fonts.googleapis.com/css2?family=Tajawal&display=swap"
      rel="stylesheet">
\"\"\"
st.markdown(RTL_CSS, unsafe_allow_html=True)
st.title("🎓 لوحة تحكم مستشار الجامعة")""")

h('A.4.2 – Dynamic File Loader for CSV and Excel Uploads', 3)
code("""import pandas as pd
import streamlit as st

@st.cache_data(show_spinner="جاري تحميل البيانات...")
def load_file(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith('.csv'):
        # Try UTF-8-SIG first (handles Arabic BOM), fall back to cp1256
        try:
            return pd.read_csv(uploaded_file, encoding='utf-8-sig')
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, encoding='cp1256')
    elif name.endswith(('.xlsx', '.xls')):
        return pd.read_excel(uploaded_file, engine='openpyxl')
    else:
        st.error("Unsupported file format. Please upload CSV or Excel.")
        return None

uploaded = st.sidebar.file_uploader(
    "📂 رفع ملف البيانات",
    type=['csv', 'xlsx', 'xls'],
    help="ادعم ملفات CSV و Excel بترميز UTF-8 أو Windows-1256"
)

if uploaded:
    df = load_file(uploaded)
    if df is not None:
        st.success(f"✅ تم تحميل {len(df):,} سجل، {len(df.columns)} عمود")
        st.dataframe(df.head(10), use_container_width=True)
else:
    st.info("الرجاء رفع ملف بيانات للبدء.")
    st.stop()""")

h('A.4.3 – Interactive Filter System for Gender-Based Exploration', 3)
code("""# Assumes df has a column named 'الجنس' (Gender in Arabic)
GENDER_COL = 'الجنس'

st.sidebar.header("🔍 فلترة البيانات")

if GENDER_COL in df.columns:
    gender_options = ['الكل'] + sorted(df[GENDER_COL].dropna().unique().tolist())
    selected_gender = st.sidebar.selectbox("الجنس", gender_options)

    if selected_gender != 'الكل':
        filtered_df = df[df[GENDER_COL] == selected_gender]
    else:
        filtered_df = df.copy()
else:
    st.sidebar.warning(f"العمود '{GENDER_COL}' غير موجود في الملف.")
    filtered_df = df.copy()

# Optional: age / year range filter
if 'السنة' in df.columns:
    years = sorted(df['السنة'].dropna().unique().tolist())
    selected_years = st.sidebar.multiselect(
        "السنوات", options=years, default=years)
    filtered_df = filtered_df[filtered_df['السنة'].isin(selected_years)]

st.sidebar.markdown(f"**السجلات المعروضة:** {len(filtered_df):,}")""")

h('A.4.4 – Treemap Visualization for Education and Employment', 3)
code("""import plotly.express as px

# Expected columns: المنطقة (Region), التخصص (Major), مستوى التوظيف (Employment)
PATH_COLS  = ['المنطقة', 'التخصص', 'مستوى التوظيف']
VALUE_COL  = 'عدد الطلاب'

missing_cols = [c for c in PATH_COLS + [VALUE_COL] if c not in filtered_df.columns]
if missing_cols:
    st.warning(f"الأعمدة التالية غير موجودة: {missing_cols}")
else:
    st.subheader("🗺️ خريطة شجرية: المنطقة > التخصص > مستوى التوظيف")

    tree_df = (
        filtered_df
        .dropna(subset=PATH_COLS + [VALUE_COL])
        .groupby(PATH_COLS, as_index=False)[VALUE_COL].sum()
    )

    fig = px.treemap(
        tree_df,
        path   = PATH_COLS,
        values = VALUE_COL,
        color  = VALUE_COL,
        color_continuous_scale = 'Blues',
        title  = "توزيع الطلاب حسب المنطقة والتخصص ومستوى التوظيف",
    )
    fig.update_layout(
        font_family = 'Tajawal',
        margin      = dict(t=50, l=10, r=10, b=10),
    )
    st.plotly_chart(fig, use_container_width=True)""")

h('A.4.5 – Export Filtered Data as CSV (Arabic Compatible)', 3)
code("""from io import BytesIO, StringIO
import pandas as pd
import streamlit as st

def to_csv_arabic(df: pd.DataFrame) -> bytes:
    \"\"\"Export DataFrame to UTF-8-SIG encoded CSV (Excel-compatible Arabic).\"\"\"
    buffer = StringIO()
    df.to_csv(buffer, index=False, encoding='utf-8-sig')
    return buffer.getvalue().encode('utf-8-sig')

def to_excel_arabic(df: pd.DataFrame) -> bytes:
    \"\"\"Export DataFrame to Excel with Arabic-compatible formatting.\"\"\"
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='البيانات المصفاة')
        ws = writer.sheets['البيانات المصفاة']
        # Auto-fit column widths
        for col in ws.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    return output.getvalue()

st.subheader("📥 تصدير البيانات")
col1, col2 = st.columns(2)

with col1:
    st.download_button(
        label    = "⬇️ تنزيل CSV (عربي)",
        data     = to_csv_arabic(filtered_df),
        file_name = "university_advisor_filtered.csv",
        mime     = "text/csv",
    )

with col2:
    st.download_button(
        label    = "⬇️ تنزيل Excel",
        data     = to_excel_arabic(filtered_df),
        file_name = "university_advisor_filtered.xlsx",
        mime     = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

st.caption(f"سيتم تصدير {len(filtered_df):,} سجل.")""")

doc.save(out)
print(f'Chapter 9 (complete) saved to {out}')
