import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_path = r'C:\Users\ALAOI\university_advisor\Chapter_7_8_English.docx'
doc = docx.Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def b(text):
    doc.add_paragraph(text, style='List Bullet')

# ══════════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════════
h('Chapter 7', 1)
h('Conclusion & Future Work', 2)
p("This chapter synthesizes the outcomes of the University Advisor System project by presenting conclusions for each of the four core components — Front-End, Back-End, AI, and Dashboard — and outlining a concrete roadmap of future enhancements that can extend the system's capabilities and impact. The conclusions reflect on what was achieved relative to the goals set at the outset, while the future work sections identify the most promising directions for continued development.")

# ── 7.1 Front-End ─────────────────────────────────────
h('7.1 Front-End Section', 2)

h('7.1.1 Conclusion', 3)
p("The front-end of the University Advisor System was successfully implemented as a modern, responsive, and role-aware user interface that meets and in several areas exceeds the requirements defined in Chapter 2. The adoption of Tailwind CSS as the styling framework proved to be one of the most impactful technical decisions of the project. Its utility-first approach enabled the team to rapidly prototype and iterate on the UI without the constraints of pre-designed components, resulting in a distinctive and professional visual identity for the platform.")
p("The use of HTML5 semantic elements and Django's template inheritance system produced a codebase that is not only functionally correct but also highly maintainable. The master base.html template ensured that changes to the global layout — such as updating the navigation structure or footer content — were immediately reflected across all pages without requiring individual file modifications.")
p("JavaScript was used judiciously, applied only where client-side interactivity genuinely improved the user experience: real-time form validation, asynchronous API calls for AI classification results, and dynamic menu behavior. This selective approach kept the client-side code lean and free of unnecessary dependencies.")
p("The role-based access control system at the template level — implemented using Django's server-side template tags — proved to be a more secure and reliable approach than client-side conditional rendering, ensuring that unauthorized UI elements were never delivered to the browser in the first place. All 16 front-end functional and role-based test cases passed, and the Lighthouse performance audit returned a score of 91/100, affirming the quality and efficiency of the implementation.")

h('7.1.2 Future Work', 3)
p("While the current front-end implementation is complete and fully functional, several enhancements are identified to further elevate the user experience in future iterations:")
b("Progressive Web App (PWA): Convert the front-end into a Progressive Web App by adding a service worker and Web App Manifest. This would enable offline access to previously visited pages, push notifications for article status updates, and installability on mobile devices — significantly increasing engagement for student and author users.")
b("Dark Mode Support: Extend the Tailwind configuration to include a dark mode color palette. A user-controlled theme toggle, with the preference persisted in localStorage, would improve accessibility and cater to users who prefer low-light interfaces.")
b("Internationalization (i18n): Integrate Django's internationalization framework to provide multi-language support, beginning with Arabic. Given the system's target deployment in Saudi Arabian universities, native Arabic language support is a high-priority enhancement.")
b("Component-Based Architecture with a JS Framework: For future major version updates, migrating the front-end to a JavaScript framework such as Vue.js or React would enable a richer, more interactive single-page application experience, particularly for complex views like the editor review panel and the real-time dashboard.")
b("Accessibility Audit and WCAG AAA Compliance: Conduct a full WCAG 2.1 Level AAA accessibility audit and implement all recommended improvements, including enhanced keyboard navigation, screen reader optimization, and focus management for modal dialogs.")

# ── 7.2 Back-End ──────────────────────────────────────
h('7.2 Back-End Section', 2)

h('7.2.1 Conclusion', 3)
p("The back-end of the University Advisor System was implemented as a secure, well-structured, and fully featured Django application backed by a PostgreSQL relational database. The decision to adopt Django's 'batteries-included' philosophy paid significant dividends throughout the development lifecycle: the built-in ORM eliminated the need for raw SQL in most cases, the admin panel provided an immediate and powerful management interface, the migration system tracked all schema changes with precision, and the authentication framework provided production-grade security with minimal configuration.")
p("The custom User model, implemented by extending AbstractUser from the outset, provided the flexibility needed to support the system's four distinct user roles without any database restructuring. The Author and Editor extension models, linked via OneToOneField, followed the established Django pattern for profile-based user extension and proved to be a clean, maintainable solution.")
p("The role-based access control system, implemented through custom decorators applied at the view level, ensured that business rules were enforced at the server — the most secure layer of the application. All 10 automated component-level tests passed, and the system demonstrated compatibility across all tested development and staging environments. The complete fulfillment of all eight original back-end requirements, as documented in Chapter 6, confirms that the back-end implementation fully achieved its goals.")

h('7.2.2 Future Work', 3)
p("The following enhancements are recommended to evolve the back-end in future iterations:")
b("Full REST API with Django REST Framework (DRF): Refactor the existing JSON views into a formally structured REST API using Django REST Framework, with proper serializers, viewsets, and OpenAPI/Swagger documentation. This would lay the foundation for a future mobile application and third-party integrations.")
b("Redis Caching Layer: Introduce Redis as a caching layer for frequently accessed, rarely changing data — such as the list of published articles and journal details. This would significantly reduce database load under high-traffic conditions.")
b("Asynchronous Task Queue with Celery: Integrate Celery with a Redis or RabbitMQ broker to handle time-intensive operations — such as sending email notifications, generating PDF exports, and triggering AI model inference — asynchronously in the background, preventing UI blocking.")
b("Full-Text Search with PostgreSQL: Implement full-text search on the article title, abstract, and keyword fields using PostgreSQL's built-in tsvector and GIN indexing. This would provide a dramatically more powerful and relevant search experience for journal readers.")
b("API Rate Limiting and Throttling: Implement request rate limiting on all API endpoints to prevent abuse and ensure fair resource allocation among users, which is critical for a publicly accessible academic platform.")

# ── 7.3 AI ────────────────────────────────────────────
h('7.3 AI Section', 2)

h('7.3.1 Conclusion', 3)
p("The AI component of the University Advisor System successfully delivers two intelligent features: a deep learning-based image classification module and an AI-powered conversational advisor chatbot. Both features are fully integrated into the platform and accessible to users through a clean and intuitive interface.")
p("The image classification model, built on a pre-trained ResNet50 backbone using PyTorch and fine-tuned on the project's target dataset, achieved a test accuracy of 91.4% and a macro-average F1-score of 0.912. These results represent strong performance for a transfer learning approach and validate the effectiveness of the two-phase training strategy — initial head training to converge the custom classification head, followed by full fine-tuning to adapt the pre-trained backbone features to the target domain.")
p("The data augmentation pipeline played a critical role in the model's ability to generalize, reducing the train-validation accuracy gap from over 8% (without augmentation) to less than 2.4% (with augmentation). The Cosine Annealing learning rate scheduler produced stable and smooth convergence during the fine-tuning phase, contributing to the final performance level.")
p("The conversational chatbot, powered by the Google Gemini API, provides students with natural language academic advising — answering questions about majors, career paths, and course requirements in a contextually aware manner. This feature transforms the platform from a passive information repository into an active, intelligent advising partner.")
p("The most significant limitation identified is the model's behavior on out-of-distribution inputs — images that do not belong to any of the target classes. As a closed-set softmax classifier, the model will always assign a confident prediction to any input, which can produce misleading results for irrelevant images. This is identified as the primary area for future improvement.")

h('7.3.2 Future Work', 3)
p("The following research and engineering directions are proposed for the AI component:")
b("Out-of-Distribution (OOD) Detection: Implement an energy-based or Mahalanobis distance-based OOD detection method that identifies inputs not belonging to any known class and returns a 'No Match Found' response rather than a potentially misleading confident prediction.")
b("Model Calibration with Temperature Scaling: Apply post-training temperature scaling to calibrate the model's confidence scores, ensuring that a 90% confidence prediction corresponds to approximately 90% actual accuracy. This is critical for responsible AI deployment in an academic context.")
b("Larger and More Diverse Dataset: Systematically expand the training dataset by collecting additional images from academic databases, applying web scraping from licensed sources, and using synthetic data generation techniques (e.g., DALL-E, Stable Diffusion) to augment under-represented classes.")
b("Explainability with Grad-CAM: Integrate Gradient-weighted Class Activation Mapping (Grad-CAM) to generate visual heatmaps overlaid on input images, highlighting the regions of the image that most influenced the model's prediction. This would enhance user trust and provide valuable diagnostic information.")
b("Chatbot Memory and Personalization: Enhance the Gemini-powered chatbot to maintain conversation history and user context across sessions, enabling personalized advising that adapts to the student's academic history, test results, and stated interests over time.")
b("Recommendation Engine Integration: Develop a collaborative filtering or content-based recommendation engine that analyzes aptitude test results across the platform's entire user base to provide peer-informed major recommendations — 'Students with similar profiles chose these majors.'")

# ── 7.4 Dashboard ─────────────────────────────────────
h('7.4 Dashboard Section', 2)

h('7.4.1 Conclusion', 3)
p("The administrative dashboard was successfully implemented as an interactive, data-driven analytics hub using the Streamlit framework, Pandas for data processing, and Plotly for visualization. The dashboard delivers on its core mandate: providing administrators with real-time, actionable insights into platform activity, user engagement, and editorial workflow metrics through an intuitive interface that requires no technical expertise to operate.")
p("The benchmark evaluation demonstrated that the dashboard performs excellently within the expected operational data volume, with cached load times under 0.5 seconds for up to 10,000 records. The @st.cache_data decorator with a 5-minute TTL was identified as the single most impactful optimization, reducing database query overhead by over 90% under normal operating conditions.")
p("The user experience evaluation confirmed that administrators could complete all core analytical tasks — viewing KPIs, applying filters, drilling down into the treemap, and exporting reports — within target completion times and without requiring technical guidance. The Plotly-based interactive charts proved particularly valuable, enabling administrators to explore data dimensions without requiring any additional query or report generation.")
p("The technology stack selection — Streamlit, Pandas, Plotly, and Openpyxl — proved to be highly synergistic. All four libraries are Python-native, eliminating the need for a separate front-end development effort for the dashboard and allowing the same team members who developed the back-end to build and maintain the analytics layer. The Excel export functionality provided by Openpyxl bridges the gap between the digital dashboard and traditional administrative reporting workflows, ensuring broad usability across all administrator profiles.")

h('7.4.2 Future Work', 3)
p("The following enhancements are proposed to extend the dashboard's capabilities in future development cycles:")
b("Real-Time Data Streaming: Replace the 5-minute cache TTL with a WebSocket-based real-time data feed using Streamlit's st.empty() and a background thread, enabling live-updating KPI cards and charts that reflect database changes without any manual refresh.")
b("Role-Specific Dashboard Views: Extend dashboard access beyond administrators to include Editor-specific views — showing their personal review queue metrics, turnaround times, and acceptance rates — and Author-specific views showing their submission history and performance analytics.")
b("Predictive Analytics: Integrate a time-series forecasting model (e.g., Facebook Prophet or ARIMA) to predict future article submission volumes and user registration trends, enabling administrators to proactively plan review capacity and platform resources.")
b("Scheduled Automated Reports: Implement a Celery-based scheduled task that automatically generates a weekly PDF summary report using ReportLab or WeasyPrint and distributes it to all administrators via email, reducing the burden of manual reporting.")
b("Advanced Filtering and Saved Views: Allow administrators to save frequently used filter combinations as named 'views' that can be quickly recalled, and extend filtering capabilities to include multi-dimensional filters such as author institution, article topic, and reviewer assignment.")
b("PostgreSQL Query Optimization: Add composite indexes on the most frequently filtered columns (status, submission_date, region, category) and implement materialized views for the most computationally expensive aggregations to maintain sub-second query times as the dataset grows to hundreds of thousands of records.")

# ══════════════════════════════════════════════════════
# CHAPTER 8 – REFERENCES
# ══════════════════════════════════════════════════════
h('Chapter 8', 1)
h('References', 2)

references = [
    "[1] Django Software Foundation. (2023). Django Documentation (Version 4.2 LTS). Retrieved from https://docs.djangoproject.com/en/4.2/",
    "[2] PostgreSQL Global Development Group. (2023). PostgreSQL 15 Documentation. Retrieved from https://www.postgresql.org/docs/15/",
    "[3] Tailwind Labs. (2023). Tailwind CSS Documentation (Version 3.x). Retrieved from https://tailwindcss.com/docs",
    "[4] PyTorch Contributors. (2023). PyTorch Documentation (Version 2.0). Retrieved from https://pytorch.org/docs/stable/index.html",
    "[5] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 770–778. https://doi.org/10.1109/CVPR.2016.90",
    "[6] Tan, M., & Le, Q. V. (2019). EfficientNet: Rethinking model scaling for convolutional neural networks. In Proceedings of the 36th International Conference on Machine Learning (ICML), 97, 6105–6114.",
    "[7] Shorten, C., & Khoshgoftaar, T. M. (2019). A survey on image data augmentation for deep learning. Journal of Big Data, 6(1), 1–48. https://doi.org/10.1186/s40537-019-0197-0",
    "[8] Loshchilov, I., & Hutter, F. (2017). SGDR: Stochastic gradient descent with warm restarts. International Conference on Learning Representations (ICLR). https://arxiv.org/abs/1608.03983",
    "[9] Nielsen, J. (1994). Usability Engineering. Academic Press.",
    "[10] Streamlit Inc. (2023). Streamlit Documentation. Retrieved from https://docs.streamlit.io/",
    "[11] The Pandas Development Team. (2023). pandas: Powerful Python data analysis toolkit (Version 2.0). https://doi.org/10.5281/zenodo.7941880",
    "[12] Plotly Technologies Inc. (2023). Collaborative data science: Plotly Python graphing library. Retrieved from https://plotly.com/python/",
    "[13] Clark, P., Openpyxl contributors. (2023). openpyxl – A Python library to read/write Excel 2010 xlsx/xlsm files. Retrieved from https://openpyxl.readthedocs.io/",
    "[14] Mozilla Developer Network (MDN). (2023). HTML: HyperText Markup Language. Retrieved from https://developer.mozilla.org/en-US/docs/Web/HTML",
    "[15] ECMA International. (2023). ECMAScript 2023 Language Specification (ECMA-262). Retrieved from https://www.ecma-international.org/publications-and-standards/standards/ecma-262/",
    "[16] Google. (2023). Gemini API Documentation. Retrieved from https://ai.google.dev/docs",
    "[17] Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On calibration of modern neural networks. Proceedings of the 34th International Conference on Machine Learning (ICML), 70, 1321–1330.",
    "[18] Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. In Proceedings of the IEEE International Conference on Computer Vision (ICCV), 618–626.",
    "[19] Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural Computation, 9(8), 1735–1780.",
    "[20] Fielding, R. T. (2000). Architectural styles and the design of network-based software architectures (Doctoral dissertation). University of California, Irvine.",
    "[21] Deng, J., Dong, W., Socher, R., Li, L.-J., Li, K., & Fei-Fei, L. (2009). ImageNet: A large-scale hierarchical image database. In Proceedings of IEEE CVPR, 248–255.",
    "[22] Chollet, F. (2021). Deep Learning with Python (2nd ed.). Manning Publications.",
    "[23] Grinberg, M. (2018). Flask Web Development: Developing Web Applications with Python (2nd ed.). O'Reilly Media.",
    "[24] Percival, H. (2017). Test-Driven Development with Python: Obey the Testing Goat (2nd ed.). O'Reilly Media.",
    "[25] W3C. (2018). Web Content Accessibility Guidelines (WCAG) 2.1. Retrieved from https://www.w3.org/TR/WCAG21/",
]

for ref in references:
    p(ref)

doc.save(output_path)
print(f'Document saved to {output_path}')
