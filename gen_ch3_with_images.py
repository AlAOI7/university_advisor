import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

output_path = r'C:\Users\ALAOI\university_advisor\Chapter_3_WithImages.docx'
doc = docx.Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def b(text):
    doc.add_paragraph(text, style='List Bullet')

def img_placeholder(figure_number, caption):
    """Add a bordered image placeholder with figure name centered."""
    # Empty lines for image space
    for _ in range(8):
        doc.add_paragraph()
    # Dashed border label paragraph
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run(f'[ Figure {figure_number}: {caption} ]')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    # Bottom space
    doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, v in enumerate(headers):
        hdr[i].text = v
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            row[ci].text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
h('Chapter 3', 1)
h('System Design', 2)
p("This chapter provides a comprehensive presentation of the University Advisor System's design across all four major components: Front-End, Back-End, AI, and Dashboard. Each section describes the architectural decisions, interface designs, data models, and intelligent components that together form a cohesive, scalable, and user-friendly academic advising platform.")

# ══════════════════════════════════════════════════════
# 3.1 Front-End Section
# ══════════════════════════════════════════════════════
h('3.1 Front-End Section', 2)
p("The front-end of the University Advisor System is engineered to deliver an intuitive, modern, and fully responsive user experience. The interface is built with HTML5, CSS3, JavaScript, and Tailwind CSS, and adapts seamlessly to all screen sizes and device types.")

# 3.1.1
h('3.1.1 Homepage – Main Entry Point', 3)
p("The Homepage serves as the primary landing page for all users. It presents a compelling introduction to the platform, featuring a hero section with a clear call-to-action guiding visitors to register or log in. It highlights the system's core features: the academic advising engine, the scientific journal, and the AI-powered classification tool. The page is fully responsive across desktop, tablet, and mobile.")
img_placeholder("3.1", "Homepage – Main Entry Point")

# 3.1.2
h('3.1.2 Scientific Journal – Main Page', 3)
p("The Scientific Journal Main Page displays a structured list of published journal issues and articles, each showing its title, abstract snippet, author names, keywords, and publication date. Users can browse by issue, filter by keyword or category, and search for specific topics. Pagination controls allow efficient navigation through large volumes of content.")
img_placeholder("3.2", "Scientific Journal – Main Page")

# 3.1.3
h('3.1.3 Authentication – Login View', 3)
p("The Authentication Login View provides a secure interface for users to access their accounts. The form requests the user's registered email and password. Client-side validation provides immediate feedback, while server-side authentication verifies credentials before granting access. Upon successful login, users are redirected to their role-specific dashboard.")
img_placeholder("3.3", "Authentication – Login View")

# 3.1.4
h('3.1.4 Scientific Journal – Article Details Page', 3)
p("The Article Details Page presents a comprehensive view of a single research article, displaying the full title, authors and their affiliations, publication date and issue number, full abstract, keywords, and a PDF download link where available. The page also lists the article's references in standardized academic format and suggests related articles.")
img_placeholder("3.4", "Scientific Journal – Article Details Page")

# 3.1.5
h('3.1.5 Editor Access – Dashboard Visibility', 3)
p("This view describes the portal editors see upon logging in. Based on the editor's role and assigned permissions, the navigation menu dynamically updates to display editor-specific options such as the Article Review Panel and Issue Management. The main surface provides a summary of pending articles, recent submissions, and notifications.")
img_placeholder("3.5", "Editor Access – Dashboard Visibility")

# 3.1.6
h('3.1.6 Editor Dashboard – Article Review Panel', 3)
p("The Article Review Panel is a dedicated workspace for editors to manage the peer-review process. It presents a tabular list of all articles assigned to the editor, categorized by review status: Pending, Under Review, Revision Requested, Accepted, or Rejected. Editors can filter and sort articles and take actions directly from the panel.")
img_placeholder("3.6", "Editor Dashboard – Article Review Panel")

# 3.1.7
h('3.1.7 Editor Dashboard – Article Review Interface', 3)
p("The Article Review Interface provides editors with all information and tools needed to evaluate a manuscript. Full article details are displayed alongside a dedicated review form where the editor records their evaluation, selects a recommendation (Accept, Minor Revision, Major Revision, Reject), and writes detailed comments. The submitted review is automatically timestamped and the author is notified.")
img_placeholder("3.7", "Editor Dashboard – Article Review Interface")

# 3.1.8
h('3.1.8 Author Access – Dashboard Visibility', 3)
p("The Author's dashboard is tailored to the needs of a researcher submitting work to the journal. Upon login, an author is greeted with a personalized dashboard showing a summary of submitted articles and their current statuses. Navigation provides access to the article submission form, profile, and notifications from editors.")
img_placeholder("3.8", "Author Access – Dashboard Visibility")

# 3.1.9
h('3.1.9 Author Dashboard – Article Management Interface', 3)
p("The Author's Article Management Interface gives researchers a complete overview of all their submitted manuscripts. Each article is listed with its title, submission date, current review status, and any editor feedback. Authors can click to view full details, read editor comments, download review reports, and submit revised versions when requested.")
img_placeholder("3.9", "Author Dashboard – Article Management Interface")

# 3.1.10
h('3.1.10 Author Dashboard – Submit Article Form', 3)
p("The Submit Article Form guides authors through the manuscript submission process in a logical, step-by-step format. Authors enter the article title, abstract, list all co-authors, add keywords, specify the research domain, and upload the manuscript file. Optional supplementary material fields are also provided. Upon submission, a confirmation notification is sent to the author.")
img_placeholder("3.10", "Author Dashboard – Submit Article Form")

# 3.1.11
h('3.1.11 User Profile – Account Information View', 3)
p("The User Profile page allows all registered users to view and manage their personal account information including display name, email, affiliated institution, and contact details. A profile picture upload option is provided. Authors can additionally manage their professional biography, research interests, and ORCID identifier. A password change section is also available.")
img_placeholder("3.11", "User Profile – Account Information View")

# 3.1.12
h('3.1.12 AI Classification – Landing & Instructions Interface', 3)
p("The AI Classification Landing Interface serves as the entry point to the system's AI-powered image classification module. This page clearly explains the tool's purpose and provides step-by-step instructions: upload an image, initiate the prediction, then view the results. The interface is designed to be approachable for non-technical users.")
img_placeholder("3.12", "AI Classification – Landing & Instructions Interface")

# 3.1.13
h('3.1.13 AI Classification – Image Upload Interface', 3)
p("The Image Upload Interface features a clear drag-and-drop area for uploading images, with a fallback 'Browse Files' button. The interface displays an immediate preview of the uploaded image. Accepted file formats (JPG, PNG) and maximum file size limits are indicated. A 'Classify' button submits the image to the AI model for processing.")
img_placeholder("3.13", "AI Classification – Image Upload Interface")

# 3.1.14
h('3.1.14 AI Classification – Prediction Result Interface', 3)
p("The Prediction Result Interface displays the output of the AI classification model. The top predicted class is shown prominently with a confidence score as a percentage. A bar chart shows the probability distribution across all categories. The original uploaded image is displayed alongside the results and users can choose to classify a new image or save the results.")
img_placeholder("3.14", "AI Classification – Prediction Result Interface")

# 3.1.15
h('3.1.15 Dashboard – Sales by Category', 3)
p("The 'Sales by Category' dashboard view provides administrators with a visual breakdown of platform engagement metrics organized by category — such as article submissions, user registrations, or test completions segmented by academic field. An interactive bar or pie chart allows quick identification of the most active categories, with date range and category filters for detailed analysis.")
img_placeholder("3.15", "Dashboard – Sales by Category Chart")

# 3.1.16
h('3.1.16 Dashboard – Treemap by Region > Category > Sub-Category', 3)
p("The Treemap Dashboard View offers a hierarchical visualization of platform data organized from Region down to Category and Sub-Category. Larger rectangles represent higher volumes of activity. Administrators can click any segment to drill down further. This provides a powerful multi-dimensional view for high-level strategic analysis.")
img_placeholder("3.16", "Dashboard – Treemap by Region > Category > Sub-Category")

# ══════════════════════════════════════════════════════
# 3.2 Back-End Section
# ══════════════════════════════════════════════════════
h('3.2 Back-End Section', 2)
p("The back-end architecture is built on the Django web framework using Python, adhering to the Model-View-Template (MVT) design pattern. PostgreSQL serves as the primary relational database, ensuring data integrity, reliability, and support for complex queries.")

h('3.2.2 Model Details', 3)
p("The system's database schema is composed of a set of interconnected models representing the core entities of the platform. Each model is implemented as a Django ORM class mapping directly to a table in the PostgreSQL database.")

h('3.2.2.1 User Model', 4)
p("The User Model is the foundational entity extending Django's AbstractUser. It stores authentication credentials (username, email, hashed password) and profile fields including full name, affiliated institution, and role (Student, Author, Editor, Administrator). The role field is critical for enforcing role-based access control throughout the application.")

h('3.2.2.2 Journal Model', 4)
p("The Journal Model represents a scientific journal hosted on the platform. It contains the journal's title, ISSN, description, founding date, editor-in-chief (ForeignKey to User), and current status (Active/Archived). It serves as the top-level container for organizing all published issues and articles.")

h('3.2.2.3 Issue Model', 4)
p("The Issue Model represents a specific volume and issue of a journal, related to the Journal Model via ForeignKey (one journal → many issues). Key attributes include volume number, issue number, publication date, and issue description. It serves as the organizational unit grouping articles published together in a single release.")

h('3.2.2.4 Author Model', 4)
p("The Author Model maintains a OneToOne relationship with the User Model and stores additional academic details: professional biography, ORCID identifier, primary research interests, and institutional affiliations. This dedicated model keeps author-specific data separate from general user data.")

h('3.2.2.5 Editor Model', 4)
p("The Editor Model maintains a OneToOne relationship with the User Model and stores editor-specific data including area of expertise (used for intelligent article assignment) and the journals/issues the editor is responsible for managing. It is central to the workflow routing logic in the back-end.")

h('3.2.2.6 Article Model', 4)
p("The Article Model is one of the most central entities, storing manuscript title, abstract, uploaded file, submission date, and associated journal issue (ForeignKey to Issue). It maintains a ManyToMany relationship with Author and includes a status field tracking the article through the editorial workflow (Submitted → Under Review → Accepted → Published → Rejected).")

h('3.2.2.7 Review Model', 4)
p("The Review Model captures editorial review outcomes. It holds ForeignKeys to both Article and Editor, stores the editor's written comments, their overall recommendation, and the review timestamp. It forms the backbone of the peer-review audit trail, ensuring all editorial decisions are recorded and traceable.")

h('3.2.2.8 Keyword Model', 4)
p("The Keyword Model stores individual academic keywords as unique string entries. Keywords are associated with articles through the ArticleKeywords junction model, enabling efficient search and filtering by topic. Using a dedicated model ensures keywords are standardized and avoids duplication.")

h('3.2.2.9 ArticleKeywords Model', 4)
p("The ArticleKeywords Model is a junction table implementing the ManyToMany relationship between Article and Keyword. Each record links one article to one keyword, allowing an article to have multiple keywords and a keyword to be associated with multiple articles.")

h('3.2.2.10 Reference Model', 4)
p("The Reference Model stores bibliographic references cited within an article, with a ForeignKey to Article. Fields include the full citation text (APA/IEEE format), DOI if available, and the reference order. This enables the system to display well-formatted reference lists on each article's detail page.")

h('3.2.3 Key Points', 3)
p("The back-end architecture is guided by core design principles:")
b("Separation of Concerns: Business logic is encapsulated within Django's service layer and model methods, keeping views thin and focused on request/response handling.")
b("Security First: All endpoints are protected by Django's CSRF middleware, and authentication is enforced on all private routes.")
b("RESTful API Design: API endpoints follow REST conventions using appropriate HTTP methods and standardized JSON responses.")
b("Scalability: PostgreSQL and a normalized schema ensure the system handles growing volumes of users and articles without performance degradation.")
b("Data Integrity: ForeignKey and ManyToMany relationships are enforced at the database level, maintaining referential integrity at all times.")

h('3.2.5 Entity Relationship Diagram (ERD)', 3)
p("The Entity Relationship Diagram (ERD) provides a comprehensive visual representation of the database schema and the relationships between all models. The diagram illustrates both cardinality (one-to-one, one-to-many, many-to-many) and the nature of relationships between entities.")
p("Key relationships: User ↔ Author (OneToOne); User ↔ Editor (OneToOne); Journal → Issues (OneToMany); Issue → Articles (OneToMany); Article ↔ Authors (ManyToMany); Article ↔ Keywords via ArticleKeywords; Article → Reviews (OneToMany); Article → References (OneToMany).")
img_placeholder("3.17", "Entity Relationship Diagram (ERD) – University Advisor System")
img_placeholder("3.18", "ERD – Detailed Model Relationships (Zoomed View)")

# ══════════════════════════════════════════════════════
# 3.3 AI Section
# ══════════════════════════════════════════════════════
h('3.3 AI Section', 2)
p("The AI section provides intelligent, data-driven functionalities elevating the platform beyond a static repository. The AI pipeline encompasses four stages: data ingestion and preprocessing, model architecture definition, training, and evaluation.")

h('3.3.1 Data Ingestion and Preprocessing', 3)
p("Raw image datasets are ingested from designated data sources. The preprocessing pipeline performs: resizing to 224×224 pixels, pixel normalization to [0,1] range, and data augmentation (random flips, rotation, color jitter) to artificially expand the training set. The dataset is split into training, validation, and test sets using a stratified approach to ensure balanced class representation.")
img_placeholder("3.19", "AI Data Pipeline – Preprocessing and Augmentation Flow")

h('3.3.2 Model Architecture', 3)
p("The classification model is built on a Convolutional Neural Network (CNN) architecture using ResNet-50 pre-trained on ImageNet as the backbone. The pre-trained convolutional base is used as a feature extractor with weights initially frozen. A custom classification head (Global Average Pooling → Dense → Dropout → Softmax) is appended on top. The model is implemented using PyTorch.")
img_placeholder("3.20", "AI Model Architecture – ResNet-50 with Custom Classification Head")

h('3.3.3 Training Engine', 3)
p("The Adam optimizer (lr = 1e-4) and Cross-Entropy Loss are used. Training is conducted in two phases: (1) head-only training with backbone frozen (feature extraction), then (2) full fine-tuning with a lower learning rate. A Cosine Annealing scheduler gradually reduces the learning rate. Early stopping monitors validation loss to halt training when performance plateaus, preventing overfitting.")
img_placeholder("3.21", "Training Dynamics – Loss and Accuracy Curves (Both Phases)")

h('3.3.4 Evaluation and Persistence', 3)
p("Model performance is evaluated on the held-out test set using accuracy, per-class precision, recall, and F1-score. A confusion matrix identifies frequently confused classes. The best model's state dictionary is serialized to disk using PyTorch and loaded by the Django back-end at runtime to serve real-time predictions through the AI Classification API endpoint.")
img_placeholder("3.22", "Confusion Matrix – Test Set Evaluation Results")

# ══════════════════════════════════════════════════════
# 3.4 Dashboard Section
# ══════════════════════════════════════════════════════
h('3.4 Dashboard Section', 2)
p("The administrative Dashboard is the central control and analytics hub for system administrators and journal managers. Built with interactive data visualization libraries, the dashboard transforms raw database statistics into clear, actionable charts and tables.")

h('3.4.1 System Components', 3)
p("The dashboard is composed of several key components:")
b("KPI Summary Cards: At-a-glance metric cards showing Total Registered Users, Total Articles Submitted, Total Articles Published, and Pending Reviews.")
b("User Activity Charts: Line and bar charts displaying user registration trends over time and daily/monthly active users.")
b("Submission Pipeline Funnel: A funnel chart visualizing the article submission and review workflow across all stages.")
b("Sales by Category Chart: An interactive bar or pie chart breaking down platform activity by academic category.")
b("Treemap Visualization: A hierarchical treemap enabling drill-down by Region → Category → Sub-Category.")
b("Content Management Panel: Interface for administrators to manage journals, majors, and user accounts.")
b("System Logs and Audit Trail: A searchable log of recent system events for security monitoring.")

img_placeholder("3.23", "Dashboard – KPI Summary Cards and User Activity Charts")
img_placeholder("3.24", "Dashboard – Submission Pipeline Funnel Chart")
img_placeholder("3.25", "Dashboard – Treemap: Region > Category > Sub-Category")

doc.save(output_path)
print(f'Document saved to {output_path}')
