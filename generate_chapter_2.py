import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_path = r'C:\Users\ALAOI\university_advisor\Chapter_2_English.docx'
doc = docx.Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def b(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx+1].cells
        for col_idx, val in enumerate(row_data):
            row[col_idx].text = val
    doc.add_paragraph()

# ── CHAPTER 2 ──────────────────────────────────────────
h('Chapter 2', 1)
h('Requirements Analysis (Tools)', 2)
p("This chapter outlines the tools, frameworks, and technologies selected for the University Advisor System. Each tool is evaluated based on its purpose, suitability for the project's requirements, and its comparative advantage over alternatives. The analysis is organized across four primary sections: Front-End, Back-End, AI, and Dashboard.")

# ══════════════════════════════════════════════════════
# 2.1 Front-End Section
# ══════════════════════════════════════════════════════
h('2.1 Front-End Section', 2)
p("The front-end layer of the University Advisor System is responsible for all user-facing interactions. The selected tools were chosen to maximize responsiveness, accessibility, and development efficiency while producing a modern and professional interface.")

# 2.1.1 HTML
h('2.1.1 HTML: Structural Foundation', 3)
h('2.1.1.1 Purpose and Role of HTML', 4)
p("HTML5 (HyperText Markup Language) is the foundational markup language used to structure all pages within the University Advisor System. It defines the skeleton of every web page, organizing content into meaningful sections such as headers, navigation bars, article bodies, forms, and footers.")

h('2.1.1.2 Justification for Selection', 4)
p("HTML5 is the universally accepted standard for web content structure and is supported by all modern browsers without any additional dependencies. Its semantic elements (e.g., <article>, <section>, <nav>, <header>) improve code readability, maintainability, and search engine optimization (SEO). As the bedrock of web development, no alternative exists or is needed.")

h('2.1.1.3 Application Across the Platform', 4)
p("HTML5 is used across every page of the platform including the homepage, scientific journal listing, article detail pages, user authentication forms, author and editor dashboards, the AI classification interface, and the administrative dashboard. Each template is structured using semantic HTML to ensure logical content hierarchy and accessibility compliance.")

h('2.1.1.4 Reusability and Semantic Integrity', 4)
p("Django's templating engine allows HTML templates to be modular and reusable through template inheritance. A base template defines the common structure (navigation, footer, meta tags), and child templates extend it by overriding specific content blocks. This approach enforces semantic integrity across all pages while drastically reducing code duplication.")

# 2.1.2 Tailwind CSS
h('2.1.2 Tailwind CSS: Styling Framework', 3)
h('2.1.2.1 Purpose and Role of Tailwind CSS', 4)
p("Tailwind CSS is a utility-first CSS framework used to style the entire front-end of the University Advisor System. Unlike traditional CSS frameworks, Tailwind provides low-level utility classes that can be composed directly in HTML to build any design without writing custom CSS from scratch.")

h('2.1.2.2 Justification for Selection', 4)
p("Tailwind CSS was selected over Bootstrap for its superior flexibility and customizability. It does not impose pre-designed component styles, allowing the team to create a unique visual identity for the platform. Its utility-first approach results in smaller production CSS bundles (via PurgeCSS), faster styling iterations, and a consistent design system enforced through configuration.")

h('2.1.2.3 Application Across the Platform', 4)
p("Tailwind CSS is applied across all templates to handle layout (Flexbox, Grid), spacing, typography, colors, responsive breakpoints, and interactive states (hover, focus, active). It is used for building the navigation bar, hero sections, article cards, data tables, form elements, and dashboard components.")

h('2.1.2.4 Benefits in Workflow and Maintainability', 4)
p("Tailwind's configuration file (tailwind.config.js) serves as the single source of truth for the design system, defining the color palette, typography scale, and spacing units. This ensures visual consistency across the entire application. The utility-class approach also means that styling changes are localized to individual components, making the codebase easier to maintain and debug.")

# 2.1.3 JavaScript
h('2.1.3 JavaScript: Interactivity and Logic Control', 3)
h('2.1.3.1 Purpose and Role of JavaScript', 4)
p("JavaScript is the primary client-side scripting language used to add dynamic behavior and interactivity to the University Advisor System's front-end. It operates in the user's browser, enabling real-time updates, user input validation, and asynchronous communication with the server without requiring full page reloads.")

h('2.1.3.2 Justification for Selection', 4)
p("Vanilla JavaScript was chosen for its ubiquity, performance, and zero-dependency footprint. Modern JavaScript (ES6+) provides all the necessary features for the project's requirements, including Promises, async/await for handling API calls, arrow functions, and the Fetch API for asynchronous HTTP requests. Using vanilla JS avoids the overhead of loading large libraries like jQuery for tasks that modern JS handles natively.")

h('2.1.3.3 Application Across the Platform', 4)
p("JavaScript is used throughout the platform for: real-time form validation on the article submission and authentication forms; dynamic rendering of search and filter results on the journal page; asynchronous communication with the AI classification endpoint using the Fetch API; interactive chart rendering on the dashboard; and managing UI state changes such as dropdown menus, modals, and notifications.")

h('2.1.3.4 Enhancing User Experience and System Feedback', 4)
p("JavaScript significantly enhances the user experience by providing immediate feedback. Form errors are highlighted instantly without a page reload, article search results update dynamically as the user types, and AI classification results are displayed without navigating away from the upload page. Toast notifications and loading spinners are controlled via JavaScript to keep users informed of system status at all times.")

# 2.1.4 PyCharm
h('2.1.4 PyCharm: Development Environment', 3)
h('2.1.4.1 Purpose and Role of PyCharm', 4)
p("PyCharm, developed by JetBrains, is the primary Integrated Development Environment (IDE) used throughout the development of the University Advisor System. It provides a comprehensive suite of tools specifically optimized for Python and Django development, streamlining the coding workflow.")

h('2.1.4.2 Justification for Selection', 4)
p("PyCharm Professional Edition was selected for its first-class Django framework support, including automatic template tag completion, Django URL routing navigation, and an integrated database browser for PostgreSQL. Its advanced code intelligence (smart autocompletion, real-time error detection, and refactoring tools) significantly reduces bugs and accelerates development compared to general-purpose editors.")

h('2.1.4.3 Application in the Development Workflow', 4)
p("PyCharm was used for all Python and Django development tasks: writing and refactoring model, view, and URL configuration code; debugging the back-end logic using the built-in visual debugger; managing the virtual environment and project dependencies; running Django management commands from the integrated terminal; and inspecting and querying the PostgreSQL database using the built-in database tool window.")

h('2.1.4.4 Contribution to Project Quality and Maintainability', 4)
p("The IDE's real-time code analysis and PEP 8 compliance checking enforced consistent code style throughout the project. The integrated Git client facilitated version control operations, ensuring all team members' contributions were tracked and merged cleanly. PyCharm's built-in test runner was also used to execute and monitor Django unit tests, contributing to overall project quality.")

# 2.1.5 Use Case Diagram
h('2.1.5 Use Case Diagram', 3)
p("The Use Case Diagram below illustrates the interactions between the system's primary actors — Student, Author, Editor, and Administrator — and the core functionalities of the University Advisor System's front-end. Key use cases include: browsing the scientific journal, submitting articles, reviewing and approving manuscripts, taking the aptitude test, viewing major recommendations, and managing user accounts. The diagram provides a high-level overview of the system's scope from the user's perspective and was used as a primary reference during the front-end design phase.")
p("[Figure 2.1: Use Case Diagram – University Advisor System Front-End]")

# 2.1.6 Flowchart
h('2.1.6 Flowchart – Scientific Journal Section', 3)
p("The flowchart below details the user journey and process flow within the Scientific Journal section of the platform. It begins with a user visiting the journal's main page and proceeds through the paths available to different user roles. An unauthenticated visitor can browse articles and view abstracts but is prompted to log in to access full content. An authenticated Author can navigate to the submission form, complete and submit a manuscript, and then track its review status. An authenticated Editor receives a notification of a new submission, accesses the review panel, evaluates the manuscript, and records their decision, which then triggers a notification to the Author.")
p("[Figure 2.2: Flowchart – Scientific Journal Section Workflow]")

# 2.1.7 Project Schedule
h('2.1.7 Project Schedule', 3)
p("The project was executed over a structured timeline divided into key phases to ensure organized delivery and milestone tracking. The schedule below outlines the major phases, their durations, and primary deliverables.")
add_table(
    ["Phase", "Duration", "Key Activities", "Deliverables"],
    [
        ["Phase 1: Requirements & Planning", "Week 1–2", "Stakeholder analysis, requirements gathering, tool selection", "Requirements document, project plan"],
        ["Phase 2: System Design", "Week 3–4", "Database schema design, ERD, UI wireframes, use case diagrams", "Design documents, ERD, wireframes"],
        ["Phase 3: Front-End Development", "Week 5–7", "HTML/Tailwind CSS templates, JavaScript logic, responsive design", "Functional UI templates"],
        ["Phase 4: Back-End Development", "Week 6–9", "Django models, views, URL routing, API endpoints, authentication", "Working back-end with database"],
        ["Phase 5: AI Model Development", "Week 8–10", "Data preprocessing, model training (PyTorch), API integration", "Trained model, classification endpoint"],
        ["Phase 6: Dashboard Development", "Week 9–11", "Streamlit/Plotly dashboard, data visualization components", "Functional analytics dashboard"],
        ["Phase 7: Integration & Testing", "Week 11–12", "System integration, unit testing, bug fixing, performance testing", "Tested, integrated application"],
        ["Phase 8: Documentation & Delivery", "Week 13–14", "Report writing, final review, deployment preparation", "Final report, deployed application"],
    ]
)

# 2.1.8 Tool Comparisons
h('2.1.8 Tool Comparisons: Tailwind CSS vs Bootstrap and JavaScript vs jQuery', 3)
p("The following tables provide a side-by-side comparison of the selected front-end tools against their most common alternatives, justifying the final technology choices made for this project.")

p("Table 2.1: Tailwind CSS vs Bootstrap")
add_table(
    ["Criteria", "Tailwind CSS", "Bootstrap"],
    [
        ["Design Philosophy", "Utility-first; full design freedom", "Component-based; pre-designed components"],
        ["Customization", "Highly customizable via config file", "Requires overriding default styles"],
        ["Bundle Size", "Minimal (purges unused CSS)", "Larger default bundle size"],
        ["Learning Curve", "Moderate (requires utility class knowledge)", "Low (ready-to-use components)"],
        ["Unique Design", "Easy to achieve unique designs", "Difficult to avoid 'Bootstrap look'"],
        ["Project Fit", "✅ Selected – ideal for custom UI", "Not selected"],
    ]
)

p("Table 2.2: JavaScript vs jQuery")
add_table(
    ["Criteria", "Vanilla JavaScript (ES6+)", "jQuery"],
    [
        ["Performance", "Native browser execution, faster", "Adds overhead from library loading"],
        ["Bundle Size", "Zero additional dependencies", "Adds ~87KB minified"],
        ["Modern Support", "Fetch API, async/await built-in", "$.ajax is an older abstraction"],
        ["Browser Support", "Excellent in all modern browsers", "Designed for older browser compat."],
        ["Maintenance", "No external dependency to update", "Requires jQuery version management"],
        ["Project Fit", "✅ Selected – modern, lightweight", "Not selected"],
    ]
)

# ══════════════════════════════════════════════════════
# 2.2 Back-End Section
# ══════════════════════════════════════════════════════
h('2.2 Back-End Section', 2)
p("The back-end of the University Advisor System is built on a robust, scalable stack centered around the Django framework and PostgreSQL database. This section details the rationale behind each back-end technology choice and how it is applied within the system.")

h('2.2.1 Why Django', 3)
h('2.2.1.1 Rapid Development and Built-in Features', 4)
p("Django was selected as the back-end framework primarily for its 'batteries-included' philosophy. Out of the box, Django provides an ORM, an admin interface, URL routing, a templating engine, form handling, and session management — all of which are core requirements for this project. This dramatically reduced the development time compared to assembling these components independently with a micro-framework like Flask.")

h('2.2.1.2 Security and Community Support', 4)
p("Django is renowned for its security-first design. It provides built-in protections against the most common web vulnerabilities including SQL injection (via parameterized ORM queries), Cross-Site Scripting (XSS), Cross-Site Request Forgery (CSRF), and Clickjacking. Its large, active community ensures extensive documentation, third-party packages, and long-term support, reducing project risk.")

h('2.2.2 Database Integration with PostgreSQL', 3)
h('2.2.2.1 Configuration and Connection', 4)
p("PostgreSQL is integrated with Django using the psycopg2 database adapter. The database connection parameters (host, port, database name, user, and password) are stored securely in environment variables and referenced in Django's settings.py file, adhering to security best practices and making the configuration environment-agnostic.")

h('2.2.2.2 Efficiency in Handling Structured Data', 4)
p("PostgreSQL was chosen over SQLite for production due to its support for concurrent connections, advanced indexing, full-text search capabilities, and robust transaction management (ACID compliance). These features are critical for a multi-user platform where Authors, Editors, and Administrators may be performing concurrent read/write operations on the database.")

h('2.2.3 Model Design and Schema Implementation', 3)
h('2.2.3.1 Defining Models in Django', 4)
p("All database tables are defined as Python classes inheriting from Django's models.Model base class. Each class attribute represents a database column with a specific field type (e.g., CharField, TextField, IntegerField, ForeignKey). Django's ORM translates these model definitions into optimized SQL CREATE TABLE statements, abstracting the developer from writing raw SQL.")

h('2.2.3.2 Migrations and Schema Evolution', 4)
p("Django's migration system provides version control for the database schema. Whenever a model is modified, a new migration file is generated using the makemigrations command. The migrate command then applies these changes to the database safely. This allows the schema to evolve iteratively throughout the project lifecycle without manual database intervention.")

h('2.2.4 Back-End Logic and Business Rules', 3)
h('2.2.4.1 View-Based Logic and URL Routing', 4)
p("Business logic is implemented within Django views, which are Python functions or classes that receive an HTTP request and return an HTTP response. Django's URL dispatcher maps incoming request URLs to the appropriate view function using a clean URL configuration file (urls.py). This separation ensures that routing logic is decoupled from business logic.")

h('2.2.4.2 Role-Specific Workflows', 4)
p("The back-end implements distinct workflows for each user role. When an Author submits an article, the back-end validates the submission data, saves the article record with a 'Submitted' status, and triggers a notification to the relevant Editor. When an Editor records a review decision, the back-end updates the article's status, logs the review, and dispatches a notification email to the Author. These workflows are implemented using Django signals and custom service functions.")

h('2.2.5 Role-Based User Access', 3)
h('2.2.5.1 Role Definitions and Capabilities', 4)
add_table(
    ["Role", "Key Capabilities"],
    [
        ["Student", "Browse majors, take aptitude test, view recommendations, use AI advisor chatbot"],
        ["Author", "Submit articles, track submission status, view editor feedback, manage profile"],
        ["Editor", "View assigned submissions, conduct reviews, accept/reject/request revisions, manage issues"],
        ["Administrator", "Full system access: manage all users, journals, majors, view dashboard analytics"],
    ]
)

h('2.2.5.2 Implementation of Access Control', 4)
p("Role-Based Access Control (RBAC) is implemented using a combination of Django's built-in Groups and Permissions system and custom middleware decorators. Each view that requires a specific role is protected with a decorator (e.g., @login_required, @role_required('Editor')) that checks the authenticated user's role before allowing access. Unauthorized access attempts result in a redirect to a permission-denied page.")

h('2.2.5.3 Visual Overview of Role Permissions', 4)
p("[Figure 2.3: Role Permission Matrix – University Advisor System]")
add_table(
    ["Feature / Permission", "Student", "Author", "Editor", "Administrator"],
    [
        ["Browse Journal Articles", "✅", "✅", "✅", "✅"],
        ["Take Aptitude Test", "✅", "✅", "✅", "✅"],
        ["Submit Article", "❌", "✅", "❌", "✅"],
        ["Review Article", "❌", "❌", "✅", "✅"],
        ["Manage Users", "❌", "❌", "❌", "✅"],
        ["View Analytics Dashboard", "❌", "❌", "❌", "✅"],
        ["Use AI Chatbot", "✅", "✅", "✅", "✅"],
    ]
)

h('2.2.6 Role-Based System Behavior Diagram', 3)
p("The diagram below illustrates how the system dynamically adjusts its behavior and available features based on the authenticated user's role. At the point of login, the system queries the user's assigned role and conditionally renders role-specific navigation menus, dashboard panels, and API endpoint access. This ensures a clean, uncluttered interface tailored to each user's responsibilities.")
p("[Figure 2.4: Role-Based System Behavior Diagram]")

h('2.2.7 Django vs Flask – Comparative Analysis', 3)
add_table(
    ["Criteria", "Django", "Flask"],
    [
        ["Architecture", "Full-stack, batteries-included MVT framework", "Lightweight micro-framework"],
        ["Built-in ORM", "Yes – Django ORM with migration support", "No – requires SQLAlchemy (external)"],
        ["Admin Interface", "Yes – powerful auto-generated admin panel", "No – requires Flask-Admin (external)"],
        ["Authentication", "Yes – built-in user auth system", "No – requires Flask-Login (external)"],
        ["Security Features", "CSRF, XSS, SQL injection protection built-in", "Manual implementation required"],
        ["Scalability", "Excellent for large, complex applications", "Better for small, simple APIs"],
        ["Learning Curve", "Steeper, but very well documented", "Lower for simple projects"],
        ["Project Fit", "✅ Selected – feature-rich, secure, scalable", "Not selected – insufficient built-ins"],
    ]
)

# ══════════════════════════════════════════════════════
# 2.3 AI Section
# ══════════════════════════════════════════════════════
h('2.3 AI Section', 2)
p("The AI component of the University Advisor System integrates a deep learning-based image classification module and a conversational advisory chatbot. This section details the software and hardware requirements for developing, training, and deploying these AI capabilities.")

h('2.3.1 Software Tools', 3)
add_table(
    ["Tool / Library", "Version", "Purpose"],
    [
        ["Python", "3.10+", "Core programming language for all AI development"],
        ["PyTorch", "2.x", "Deep learning framework for model definition, training, and inference"],
        ["Torchvision", "0.15+", "Pre-trained models (ResNet, EfficientNet) and image transformation utilities"],
        ["NumPy", "1.24+", "Numerical computing and array manipulation for data processing"],
        ["Pillow (PIL)", "9.x", "Image loading, resizing, and format conversion for preprocessing pipeline"],
        ["Scikit-learn", "1.2+", "Dataset splitting, evaluation metrics (precision, recall, F1), confusion matrix"],
        ["Matplotlib / Seaborn", "3.7+", "Visualization of training curves, confusion matrix, and evaluation metrics"],
        ["Google Gemini API", "Latest", "Powers the conversational AI advisor chatbot for student guidance"],
        ["CUDA Toolkit", "11.8+", "Enables GPU-accelerated model training via NVIDIA hardware"],
    ]
)

h('2.3.2 Hardware and System Configuration', 3)
add_table(
    ["Component", "Specification", "Role in AI Development"],
    [
        ["CPU", "Intel Core i7 / AMD Ryzen 7 or higher", "General computation, data loading, preprocessing"],
        ["GPU", "NVIDIA GPU (RTX 3060 or higher recommended)", "Accelerated deep learning model training via CUDA"],
        ["RAM", "16 GB minimum (32 GB recommended)", "Holding large datasets and model parameters in memory"],
        ["Storage", "SSD with 50+ GB free space", "Fast I/O for loading training datasets during training"],
        ["OS", "Windows 10/11 or Ubuntu 20.04+", "Development and training environment"],
        ["Cloud (Optional)", "Google Colab / Kaggle Notebooks", "Free GPU access for training if local hardware is insufficient"],
    ]
)

h('2.3.3 Rationale for Tool Choices', 3)
p("PyTorch was selected as the deep learning framework due to its dynamic computation graph (define-by-run), which makes debugging and iterative model development significantly more intuitive than static graph frameworks. Its tight integration with Torchvision provides immediate access to state-of-the-art pre-trained models, enabling powerful transfer learning with minimal code. The Google Gemini API was selected for the chatbot due to its advanced natural language understanding capabilities, enabling nuanced and context-aware academic advising conversations that a custom-trained model could not realistically achieve within the project's scope.")

h('2.3.4 PyTorch vs TensorFlow – Comparative Insight', 3)
add_table(
    ["Criteria", "PyTorch", "TensorFlow / Keras"],
    [
        ["Computation Graph", "Dynamic (eager execution by default)", "Static (graph mode) / Dynamic (eager mode)"],
        ["Ease of Debugging", "Excellent – standard Python debugging tools work", "More complex – requires special TF debugger"],
        ["Research Community", "Dominant in academic research", "More widely used in industry production"],
        ["Pre-trained Models", "Torchvision, HuggingFace Hub", "TensorFlow Hub, Keras Applications"],
        ["Deployment", "TorchScript, ONNX export", "TensorFlow Serving, TensorFlow Lite"],
        ["Learning Curve", "Pythonic, easier for Python developers", "Steeper, more boilerplate in TF 1.x"],
        ["Project Fit", "✅ Selected – dynamic, Pythonic, research-grade", "Not selected"],
    ]
)

# ══════════════════════════════════════════════════════
# 2.4 Dashboard Section
# ══════════════════════════════════════════════════════
h('2.4 Dashboard Section', 2)
p("The administrative dashboard provides real-time analytics and data visualization capabilities. The following tools were selected to build a functional, interactive, and data-rich dashboard interface.")

h('2.4.1 Streamlit', 3)
p("Streamlit is the primary framework used to build the administrative analytics dashboard. It enables the creation of interactive, data-driven web applications entirely in Python, without requiring any front-end development knowledge. Streamlit automatically re-renders the dashboard whenever the underlying data or user inputs change, providing a real-time, reactive experience. Its simplicity allows back-end and data science team members to contribute to the dashboard without needing HTML or JavaScript expertise.")

h('2.4.2 Pandas', 3)
p("Pandas is the core data manipulation library used within the dashboard. It provides the DataFrame — a powerful, flexible two-dimensional data structure — which is used to load, clean, transform, filter, aggregate, and prepare data from the PostgreSQL database for visualization. All data operations in the dashboard, from calculating submission counts to aggregating user activity by region, are performed using Pandas' expressive and efficient API.")

h('2.4.3 Plotly', 3)
p("Plotly is the interactive visualization library used to render all charts and graphs within the Streamlit dashboard. Unlike static visualization libraries such as Matplotlib, Plotly generates interactive charts that support zooming, panning, hovering for data point details, and filtering by clicking on legend items. The library is used to create bar charts, line charts, pie charts, funnel charts, and the hierarchical treemap visualizations, providing administrators with an engaging and explorable data analysis experience.")

h('2.4.4 Openpyxl', 3)
p("Openpyxl is used within the dashboard to provide data export functionality. Administrators can export dashboard reports and data tables to Excel (.xlsx) format for offline analysis, record-keeping, or reporting to institutional stakeholders. Openpyxl allows the back-end to programmatically create, populate, and format Excel workbooks, including applying styles, setting column widths, and creating multiple named sheets within a single export file.")

h('2.4.4.1 Alternatives', 4)
add_table(
    ["Alternative", "Reason Not Selected"],
    [
        ["XlsxWriter", "Write-only library; cannot read existing Excel files. Less flexible for general use."],
        ["xlrd", "Read-only and limited to older .xls format. Insufficient for modern .xlsx requirements."],
        ["csv module", "Produces plain CSV files with no formatting, formulas, or multi-sheet support."],
    ]
)

h('2.4.5 Use Case Diagram', 3)
p("The Use Case Diagram for the Dashboard section illustrates the interactions between the Administrator role and the dashboard's core functionalities. Key use cases include: viewing KPI summary metrics, analyzing article submission trends, exploring the sales-by-category chart, drilling down into the regional treemap, exporting reports to Excel, and managing system data through the CMS panel. The diagram confirms that the dashboard is exclusively accessible to users with the Administrator role.")
p("[Figure 2.5: Use Case Diagram – Administrative Dashboard]")

h('2.4.6 System Workflow – Flowchart', 3)
p("The system workflow flowchart for the Dashboard section describes the data pipeline from source to visualization. The process begins with the Administrator logging in and accessing the dashboard URL. The Streamlit application initializes and establishes a connection to the PostgreSQL database via Django's ORM or a direct database connection. Pandas DataFrames are populated with the queried data. Plotly then renders the appropriate charts based on the loaded DataFrames. If the Administrator applies a filter (e.g., selecting a date range), the data query is updated and the charts re-render dynamically. The Administrator may then choose to export the current view as an Excel report using Openpyxl.")
p("[Figure 2.6: System Workflow Flowchart – Dashboard Data Pipeline]")

h('2.4.7 Streamlit vs Dash – Dashboard Framework Comparison', 3)
add_table(
    ["Criteria", "Streamlit", "Dash (by Plotly)"],
    [
        ["Primary Language", "Pure Python – no front-end code needed", "Python with React.js components"],
        ["Ease of Use", "Very high – ideal for data scientists", "Moderate – requires understanding of callbacks"],
        ["Interactivity Model", "Automatic re-run on widget change", "Explicit callback functions required"],
        ["Customization", "Moderate – limited layout control", "High – full CSS and HTML customization"],
        ["Setup Complexity", "Very low – pip install streamlit", "Low – pip install dash"],
        ["Community & Docs", "Large, fast-growing community", "Mature, well-documented community"],
        ["Best For", "Rapid prototyping, data apps, internal tools", "Complex, production-grade dashboards"],
        ["Project Fit", "✅ Selected – rapid, Pythonic, no front-end overhead", "Not selected"],
    ]
)

doc.save(output_path)
print(f'Document saved to {output_path}')
