import docx
import os

doc_path = r'C:\Users\ALAOI\university_advisor\Muthnab ScienceNet Graduation Project.2 (1) (1).docx'
doc = docx.Document(doc_path)

# Update Title
doc.paragraphs[16].text = 'نظام المرشد الجامعي\n(University Advisor System)'

# Update Authors
doc.paragraphs[19].text = 'جوري مطلق المطيري\nفجر الحميدي المطيري\nفوز فيصل العتيبي\nدانه فهد السويداني\nمنيره سليمان السدراني\nحصه محمد الحرابي'

# Update Supervisor
doc.paragraphs[22].text = 'الأستاذة: وداد العواد'

# Abstract Text (Paragraph 99)
abstract_text = 'يقدم مشروع التخرج هذا نظام "المرشد الجامعي"، وهو منصة ويب متكاملة تهدف إلى مساعدة الطلاب في اختيار التخصص الجامعي المناسب. يجمع النظام بين اختبار تفاعلي لتحديد الشخصية والميول، وقاعدة بيانات شاملة للتخصصات الجامعية، وتوصيات ذكية مبنية على نتائج الاختبار، بالإضافة إلى معلومات عن الدورات والكتب المرتبطة بكل تخصص. تم بناء الواجهة الأمامية باستخدام HTML5 و CSS3 و JavaScript مع Bootstrap 5 لتقديم تجربة مستخدم متجاوبة وعصرية. بينما تم تطوير الواجهة الخلفية باستخدام إطار عمل Django مع قواعد بيانات SQLite و MySQL، مما يوفر بنية تحتية قوية لإدارة الحسابات، التخصصات، الاختبارات، والإشعارات. بشكل عام، يوضح هذا المشروع التطبيق الفعال لتقنيات تطوير الويب الحديثة لتلبية الاحتياجات الأكاديمية الحقيقية من خلال تصميم معياري وقابل للتوسع يتمحور حول المستخدم.'

for p in doc.paragraphs:
    if 'This graduation project introduces an integrated web-based platform' in p.text:
        p.text = abstract_text

# Introduction Text
intro_texts = [
    'مع استمرار التحول الرقمي في إعادة تعريف المشهد الأكاديمي، أصبح تطوير منصات ويب متكاملة محركاً أساسياً للابتكار. يستجيب مشروع التخرج هذا لهذا التوجه من خلال تقديم نظام موحد يهدف إلى مساعدة الطلاب في اتخاذ أحد أهم القرارات في حياتهم: اختيار التخصص الجامعي.',
    'تم تقسيم عملية التطوير إلى عدة مجالات تقنية رئيسية: تطوير الواجهة الأمامية، تطوير الواجهة الخلفية، وتصميم قواعد البيانات. تطلب كل قسم تخطيطاً دقيقاً واختياراً استراتيجياً للأدوات وممارسات برمجة معيارية لضمان بقاء النظام قابلاً للتطوير وسهل الاستخدام ومحسن الأداء. مكن استخدام أطر عمل حديثة مثل Django و Bootstrap فريق العمل من بناء تطبيق قوي بهيكلية نظيفة وقابلة للصيانة وواجهة مستخدم متجاوبة.',
    'يقدم هذا الفصل نظرة عامة متعمقة على بنية المنصة ومكوناتها الفردية. يبدأ بمناقشة طبقة الواجهة الأمامية، والتي تعمل كجسر تفاعلي بين المستخدمين ووظائف النظام. يوضح القسم التالي هيكل الواجهة الخلفية، مع التركيز على التنفيذ القائم على Django والذي يدير تدفق البيانات وأدوار المستخدمين والأمان.',
    'من خلال هذا التصميم متعدد الطبقات، لا يوضح المشروع الجدوى الفنية لبناء تطبيقات ويب معقدة فحسب، بل يوضح أيضاً أهميتها العملية في السياقات الأكاديمية الواقعية.'
]

for p in doc.paragraphs:
    if 'As digital transformation continues to redefine' in p.text:
        p.text = intro_texts[0]
    elif 'The development process was divided into four major technical domains' in p.text:
        p.text = intro_texts[1]
    elif 'This chapter provides an in-depth overview' in p.text:
        p.text = intro_texts[2]
    elif 'Through this layered design' in p.text:
        p.text = intro_texts[3]

doc.save(r'C:\Users\ALAOI\university_advisor\University_Advisor_Project.docx')
print('Document updated successfully!')
