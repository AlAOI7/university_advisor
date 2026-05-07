import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

out = r'C:\Users\ALAOI\university_advisor\Chapter_9_English.docx'
doc = docx.Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

h('Chapter 9', 1)
h('Appendices', 2)
p("This chapter contains the complete source code listings for the University Advisor System, organized by component. These appendices serve as a technical reference for all major implementation artifacts.")

# ═══════════════════════════════════
h('9.1 Front-End Section', 2)
p("The following listings present the core HTML templates used in the front-end of the University Advisor System.")

h('A.1.1 – Homepage Code (index.html)', 3)
p("The homepage template extends the base layout and renders the hero section, feature highlights, and calls-to-action. It uses Tailwind CSS utility classes for layout and styling and Django template tags for dynamic content.")
code("""{% extends 'base.html' %}
{% block title %}University Advisor – Home{% endblock %}

{% block content %}
<!-- Hero Section -->
<section class="min-h-screen flex items-center justify-center bg-gradient-to-br from-blue-900 to-indigo-800">
  <div class="text-center px-6">
    <h1 class="text-5xl font-bold text-white mb-4">
      Find Your Perfect Major
    </h1>
    <p class="text-xl text-blue-200 mb-8 max-w-2xl mx-auto">
      Take our AI-powered aptitude assessment and discover the
      university major that matches your strengths and interests.
    </p>
    <div class="flex gap-4 justify-center">
      {% if user.is_authenticated %}
        <a href="{% url 'test:start' %}"
           class="px-8 py-3 bg-white text-blue-900 font-semibold
                  rounded-full hover:bg-blue-50 transition">
          Take the Test
        </a>
      {% else %}
        <a href="{% url 'accounts:register' %}"
           class="px-8 py-3 bg-white text-blue-900 font-semibold
                  rounded-full hover:bg-blue-50 transition">
          Get Started
        </a>
        <a href="{% url 'accounts:login' %}"
           class="px-8 py-3 border-2 border-white text-white
                  rounded-full hover:bg-white hover:text-blue-900 transition">
          Sign In
        </a>
      {% endif %}
    </div>
  </div>
</section>

<!-- Features Section -->
<section class="py-20 bg-gray-50">
  <div class="max-w-6xl mx-auto px-6 grid grid-cols-1 md:grid-cols-3 gap-8">
    <div class="bg-white rounded-2xl p-8 shadow-md hover:shadow-xl transition">
      <div class="text-4xl mb-4">🎓</div>
      <h2 class="text-xl font-bold text-gray-800 mb-2">
        Smart Major Finder
      </h2>
      <p class="text-gray-500">
        Our AI analyzes your aptitude test results to recommend
        the best-fit academic majors for your profile.
      </p>
    </div>
    <div class="bg-white rounded-2xl p-8 shadow-md hover:shadow-xl transition">
      <div class="text-4xl mb-4">📚</div>
      <h2 class="text-xl font-bold text-gray-800 mb-2">
        Scientific Journal
      </h2>
      <p class="text-gray-500">
        Explore peer-reviewed articles across academic disciplines
        submitted and reviewed by our expert community.
      </p>
    </div>
    <div class="bg-white rounded-2xl p-8 shadow-md hover:shadow-xl transition">
      <div class="text-4xl mb-4">🤖</div>
      <h2 class="text-xl font-bold text-gray-800 mb-2">
        AI Advisor Chat
      </h2>
      <p class="text-gray-500">
        Ask our Gemini-powered chatbot any question about majors,
        careers, or university life and get instant guidance.
      </p>
    </div>
  </div>
</section>
{% endblock %}""")

h('A.1.2 – Scientific Journal Homepage (Dynamic Article List)', 3)
p("This template renders the journal's main listing page. It uses Django's pagination and a Fetch API-powered search to dynamically filter articles by title or keyword without a full page reload.")
code("""{% extends 'base.html' %}
{% block title %}Scientific Journal – University Advisor{% endblock %}

{% block content %}
<div class="max-w-6xl mx-auto px-6 py-12">
  <h1 class="text-3xl font-bold text-gray-800 mb-2">Scientific Journal</h1>
  <p class="text-gray-500 mb-8">
    Browse peer-reviewed research articles across all disciplines.
  </p>

  <!-- Search Bar -->
  <div class="mb-8">
    <input id="search-input" type="text"
           placeholder="Search by title or keyword..."
           class="w-full md:w-1/2 px-5 py-3 border border-gray-300
                  rounded-full shadow-sm focus:outline-none
                  focus:ring-2 focus:ring-blue-500">
  </div>

  <!-- Article Grid -->
  <div id="article-grid"
       class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">
    {% for article in articles %}
    <div class="bg-white rounded-2xl p-6 shadow hover:shadow-lg transition
                article-card" data-title="{{ article.title|lower }}"
         data-keywords="{{ article.keywords_str|lower }}">
      <span class="text-xs font-semibold text-blue-600 uppercase tracking-wide">
        {{ article.category }}
      </span>
      <h2 class="text-lg font-bold text-gray-800 mt-1 mb-2 line-clamp-2">
        {{ article.title }}
      </h2>
      <p class="text-gray-500 text-sm line-clamp-3">{{ article.abstract }}</p>
      <div class="mt-4 flex items-center justify-between">
        <span class="text-xs text-gray-400">
          {{ article.submission_date|date:"M d, Y" }}
        </span>
        <a href="{% url 'journal:article_detail' article.pk %}"
           class="text-sm font-semibold text-blue-600 hover:underline">
          Read More →
        </a>
      </div>
    </div>
    {% empty %}
    <p class="col-span-3 text-center text-gray-400 py-12">
      No articles found.
    </p>
    {% endfor %}
  </div>

  <!-- Pagination -->
  {% if articles.has_other_pages %}
  <div class="flex justify-center mt-10 gap-2">
    {% if articles.has_previous %}
    <a href="?page={{ articles.previous_page_number }}"
       class="px-4 py-2 bg-white border rounded-lg hover:bg-gray-50">
      ← Prev
    </a>
    {% endif %}
    <span class="px-4 py-2 bg-blue-600 text-white rounded-lg">
      {{ articles.number }} / {{ articles.paginator.num_pages }}
    </span>
    {% if articles.has_next %}
    <a href="?page={{ articles.next_page_number }}"
       class="px-4 py-2 bg-white border rounded-lg hover:bg-gray-50">
      Next →
    </a>
    {% endif %}
  </div>
  {% endif %}
</div>

<script>
const searchInput = document.getElementById('search-input');
const cards = document.querySelectorAll('.article-card');
searchInput.addEventListener('input', () => {
  const query = searchInput.value.toLowerCase();
  cards.forEach(card => {
    const match = card.dataset.title.includes(query) ||
                  card.dataset.keywords.includes(query);
    card.style.display = match ? 'block' : 'none';
  });
});
</script>
{% endblock %}""")

h('A.1.3 – Login/Signup Page (Authentication)', 3)
p("The authentication template provides both login and registration forms in a tabbed interface. JavaScript switches between the two forms without a page reload.")
code("""{% extends 'base.html' %}
{% block title %}Sign In – University Advisor{% endblock %}

{% block content %}
<div class="min-h-screen flex items-center justify-center bg-gray-100">
  <div class="bg-white rounded-2xl shadow-lg w-full max-w-md p-8">

    <!-- Tabs -->
    <div class="flex mb-6 border-b">
      <button id="tab-login" onclick="switchTab('login')"
              class="flex-1 pb-2 font-semibold text-blue-600
                     border-b-2 border-blue-600">
        Sign In
      </button>
      <button id="tab-register" onclick="switchTab('register')"
              class="flex-1 pb-2 font-semibold text-gray-400">
        Register
      </button>
    </div>

    <!-- Login Form -->
    <form id="form-login" method="post" action="{% url 'accounts:login' %}">
      {% csrf_token %}
      <input type="hidden" name="form_type" value="login">
      <div class="mb-4">
        <label class="block text-sm font-medium text-gray-700 mb-1">
          Email
        </label>
        <input name="username" type="email" required
               class="w-full px-4 py-2 border rounded-lg
                      focus:outline-none focus:ring-2 focus:ring-blue-500">
      </div>
      <div class="mb-6">
        <label class="block text-sm font-medium text-gray-700 mb-1">
          Password
        </label>
        <input name="password" type="password" required
               class="w-full px-4 py-2 border rounded-lg
                      focus:outline-none focus:ring-2 focus:ring-blue-500">
      </div>
      {% if messages %}
        {% for message in messages %}
        <p class="text-red-500 text-sm mb-4">{{ message }}</p>
        {% endfor %}
      {% endif %}
      <button type="submit"
              class="w-full py-3 bg-blue-600 text-white font-semibold
                     rounded-lg hover:bg-blue-700 transition">
        Sign In
      </button>
      <p class="text-center text-sm text-gray-500 mt-4">
        <a href="{% url 'accounts:password_reset' %}"
           class="text-blue-600 hover:underline">
          Forgot password?
        </a>
      </p>
    </form>

    <!-- Register Form -->
    <form id="form-register" method="post" action="{% url 'accounts:register' %}"
          class="hidden">
      {% csrf_token %}
      <div class="grid grid-cols-2 gap-4 mb-4">
        <div>
          <label class="block text-sm font-medium text-gray-700 mb-1">
            First Name
          </label>
          <input name="first_name" type="text" required
                 class="w-full px-4 py-2 border rounded-lg
                        focus:outline-none focus:ring-2 focus:ring-blue-500">
        </div>
        <div>
          <label class="block text-sm font-medium text-gray-700 mb-1">
            Last Name
          </label>
          <input name="last_name" type="text" required
                 class="w-full px-4 py-2 border rounded-lg
                        focus:outline-none focus:ring-2 focus:ring-blue-500">
        </div>
      </div>
      <div class="mb-4">
        <label class="block text-sm font-medium text-gray-700 mb-1">
          Email
        </label>
        <input name="email" type="email" required
               class="w-full px-4 py-2 border rounded-lg
                      focus:outline-none focus:ring-2 focus:ring-blue-500">
      </div>
      <div class="mb-4">
        <label class="block text-sm font-medium text-gray-700 mb-1">
          Role
        </label>
        <select name="role"
                class="w-full px-4 py-2 border rounded-lg
                       focus:outline-none focus:ring-2 focus:ring-blue-500">
          <option value="student">Student</option>
          <option value="author">Author / Researcher</option>
        </select>
      </div>
      <div class="mb-6">
        <label class="block text-sm font-medium text-gray-700 mb-1">
          Password
        </label>
        <input name="password1" type="password" required
               class="w-full px-4 py-2 border rounded-lg
                      focus:outline-none focus:ring-2 focus:ring-blue-500">
      </div>
      <button type="submit"
              class="w-full py-3 bg-blue-600 text-white font-semibold
                     rounded-lg hover:bg-blue-700 transition">
        Create Account
      </button>
    </form>
  </div>
</div>

<script>
function switchTab(tab) {
  document.getElementById('form-login').classList.toggle('hidden', tab !== 'login');
  document.getElementById('form-register').classList.toggle('hidden', tab !== 'register');
  document.getElementById('tab-login').classList.toggle('text-blue-600', tab === 'login');
  document.getElementById('tab-login').classList.toggle('border-b-2', tab === 'login');
  document.getElementById('tab-register').classList.toggle('text-blue-600', tab === 'register');
  document.getElementById('tab-register').classList.toggle('border-b-2', tab === 'register');
}
</script>
{% endblock %}""")

# ═══════════════════════════════════
h('9.2 Back-End Section', 2)
p("The following listings present the core Django model implementations for the University Advisor System's back-end layer.")

h('A.2.1 – Custom User Model with Role-Based Access', 3)
p("This model extends Django's AbstractUser to add a role field and helper methods used throughout the application for permission checking.")
code("""# accounts/models.py
from django.contrib.auth.models import AbstractUser
from django.db import models

class User(AbstractUser):
    ROLE_CHOICES = [
        ('student',       'Student'),
        ('author',        'Author'),
        ('editor',        'Editor'),
        ('administrator', 'Administrator'),
    ]
    role        = models.CharField(max_length=20, choices=ROLE_CHOICES,
                                   default='student')
    institution = models.CharField(max_length=255, blank=True, null=True)
    bio         = models.TextField(blank=True, null=True)
    avatar      = models.ImageField(upload_to='avatars/', blank=True, null=True)
    created_at  = models.DateTimeField(auto_now_add=True)

    # ── Helpers ────────────────────────────────────────
    def is_student(self):
        return self.role == 'student'

    def is_author(self):
        return self.role == 'author'

    def is_editor(self):
        return self.role == 'editor'

    def is_administrator(self):
        return self.role == 'administrator'

    def get_dashboard_url(self):
        urls = {
            'student':       '/student/dashboard/',
            'author':        '/author/dashboard/',
            'editor':        '/editor/dashboard/',
            'administrator': '/admin-panel/dashboard/',
        }
        return urls.get(self.role, '/')

    def __str__(self):
        return f"{self.username} ({self.get_role_display()})"

    class Meta:
        verbose_name      = 'User'
        verbose_name_plural = 'Users'
        ordering          = ['-created_at']

# settings.py – point Django to the custom model
# AUTH_USER_MODEL = 'accounts.User'""")

h('A.2.2 – Author Profile Model', 3)
p("Extends the User model via OneToOneField to store author-specific academic metadata. A post_save signal automatically creates this profile when a User with role='author' is registered.")
code("""# journal/models.py  (Author section)
from django.db import models
from django.db.models.signals import post_save
from django.dispatch import receiver
from accounts.models import User

class Author(models.Model):
    user               = models.OneToOneField(
                             User, on_delete=models.CASCADE,
                             related_name='author_profile')
    orcid_id           = models.CharField(max_length=19, blank=True, null=True,
                             help_text="Format: 0000-0000-0000-0000")
    research_interests = models.TextField(blank=True, null=True)
    h_index            = models.PositiveIntegerField(default=0)
    total_publications = models.PositiveIntegerField(default=0)
    website            = models.URLField(blank=True, null=True)

    def get_published_articles(self):
        return self.articles.filter(status='published')

    def get_pending_articles(self):
        return self.articles.filter(
            status__in=['submitted', 'under_review', 'revision_requested'])

    def update_publication_count(self):
        self.total_publications = self.get_published_articles().count()
        self.save(update_fields=['total_publications'])

    def __str__(self):
        return f"Author: {self.user.get_full_name()}"

    class Meta:
        ordering = ['user__last_name']

# ── Auto-create Author profile on User save ────────────
@receiver(post_save, sender=User)
def create_author_profile(sender, instance, created, **kwargs):
    if created and instance.role == 'author':
        Author.objects.get_or_create(user=instance)""")

h('A.2.3 – Editor Profile Model', 3)
p("Stores editor-specific data including areas of expertise used for intelligent article assignment routing.")
code("""# journal/models.py  (Editor section)
class Editor(models.Model):
    SENIORITY_CHOICES = [
        ('junior',  'Junior Editor'),
        ('senior',  'Senior Editor'),
        ('chief',   'Editor-in-Chief'),
    ]
    user             = models.OneToOneField(
                           User, on_delete=models.CASCADE,
                           related_name='editor_profile')
    expertise_areas  = models.TextField(
                           help_text="Comma-separated list of research areas")
    seniority        = models.CharField(max_length=20,
                           choices=SENIORITY_CHOICES, default='junior')
    max_active_reviews = models.PositiveIntegerField(default=5)
    joined_editorial  = models.DateField(auto_now_add=True)

    def get_active_reviews(self):
        return self.reviews.filter(article__status='under_review')

    def is_available(self):
        return self.get_active_reviews().count() < self.max_active_reviews

    def get_expertise_list(self):
        return [e.strip() for e in self.expertise_areas.split(',')]

    def matches_article(self, article):
        article_keywords = [k.lower() for k in article.get_keyword_list()]
        return any(exp.lower() in article_keywords
                   for exp in self.get_expertise_list())

    def __str__(self):
        return f"{self.get_seniority_display()}: {self.user.get_full_name()}"

# ── Auto-create Editor profile on User save ────────────
@receiver(post_save, sender=User)
def create_editor_profile(sender, instance, created, **kwargs):
    if created and instance.role == 'editor':
        Editor.objects.get_or_create(
            user=instance,
            defaults={'expertise_areas': 'General'}
        )""")

h('A.2.4 – Article Model Reference to Author', 3)
p("The Article model is the central entity linking Authors, Editors, Keywords, and Reviews. It tracks the full editorial lifecycle via a status field.")
code("""# journal/models.py  (Article section)
class Article(models.Model):
    STATUS_CHOICES = [
        ('submitted',          'Submitted'),
        ('under_review',       'Under Review'),
        ('revision_requested', 'Revision Requested'),
        ('accepted',           'Accepted'),
        ('published',          'Published'),
        ('rejected',           'Rejected'),
    ]
    title           = models.CharField(max_length=500)
    abstract        = models.TextField()
    manuscript_file = models.FileField(upload_to='manuscripts/')
    authors         = models.ManyToManyField(Author, related_name='articles')
    issue           = models.ForeignKey('Issue', on_delete=models.SET_NULL,
                          null=True, blank=True, related_name='articles')
    status          = models.CharField(max_length=25, choices=STATUS_CHOICES,
                          default='submitted')
    submission_date = models.DateTimeField(auto_now_add=True)
    updated_at      = models.DateTimeField(auto_now=True)
    region          = models.CharField(max_length=100, blank=True)
    category        = models.CharField(max_length=100, blank=True)
    sub_category    = models.CharField(max_length=100, blank=True)

    def get_keyword_list(self):
        return list(self.keywords.values_list('keyword__name', flat=True))

    def get_primary_author(self):
        return self.authors.first()

    def assign_to_editor(self):
        available = Editor.objects.filter(
            user__is_active=True
        ).order_by('?')
        for editor in available:
            if editor.is_available() and editor.matches_article(self):
                return editor
        return available.filter(
            user__is_active=True
        ).order_by('?').first()

    def __str__(self):
        return f"[{self.get_status_display()}] {self.title[:80]}"

    class Meta:
        ordering = ['-submission_date']

# ── Signal: notify author on status change ─────────────
@receiver(post_save, sender=Article)
def notify_on_status_change(sender, instance, created, **kwargs):
    if not created:
        from .notifications import send_status_notification
        send_status_notification(instance)""")

doc.save(out)
print(f'Part 1 saved to {out}')
